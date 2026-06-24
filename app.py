import streamlit as st
import requests
import os
import io
import re
import json
import zipfile
import tempfile
import traceback
from concurrent.futures import ThreadPoolExecutor, as_completed
import xml.etree.ElementTree as ET
from copy import deepcopy
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv
import openpyxl

load_dotenv()

import sys as _sys

if getattr(_sys, "frozen", False):
    # Running as a PyInstaller exe — templates sit next to the .exe
    _TEMPLATE_BASE = os.path.dirname(_sys.executable)
else:
    # Development — templates are in the same folder as app.py (DIFY-frontend/)
    _TEMPLATE_BASE = os.path.dirname(os.path.abspath(__file__))

TEMPLATE_INDEX  = os.path.join(_TEMPLATE_BASE, "template_index.xlsx")

# Bundled MA/AR templates that ship next to the .exe (or alongside app.py in dev).
# These are the fallback used when the OneDrive-synced library isn't available.
_BUNDLED_AR_DIR = os.path.join(_TEMPLATE_BASE, "AR_template")
_BUNDLED_MA_DIR = os.path.join(_TEMPLATE_BASE, "MA_template")

# EY keeps the authoritative MA/AR templates in the SharePoint library
# "GCSOCR / Reporting files templates". Rather than ship static copies, the app can
# pull the latest .docx straight from that library at startup so templates stay
# current without rebuilding the exe. `TEMPLATE_SOURCE` selects where they come from:
#
#   bundled    (default) — use the .docx that ship next to the exe.
#   sharepoint           — download from the online SharePoint library (see below).
#   onedrive             — read from a locally synced/mapped copy at TEMPLATE_BASE_PATH.
#
# SharePoint notes: the library is gated behind each user's EY sign-in, so the app
# does NOT store any credential or use an Azure AD app registration. Instead it reuses
# the cookies of the browser the user is already signed in with on an internal machine
# (best effort, via browser_cookie3), calls the SharePoint REST API, and caches the
# .docx locally. If the fetch can't authenticate or the network is down, it falls back
# to the bundled templates and shows a warning. This path must be verified on a real
# EY-managed machine. Folder/site URLs default to the current library layout and can
# be overridden via the SHAREPOINT_* env vars.
TEMPLATE_SOURCE     = os.getenv("TEMPLATE_SOURCE", "bundled").strip().lower()

# -- onedrive (locally synced / mapped folder) mode --
TEMPLATE_BASE_PATH  = os.getenv("TEMPLATE_BASE_PATH", "").strip()
MA_TEMPLATE_SUBPATH = os.getenv("MA_TEMPLATE_SUBPATH", "1.1 MA整理版").strip()
AR_TEMPLATE_SUBPATH = os.getenv("AR_TEMPLATE_SUBPATH", "1.2 AR整理版/AR updated Verison").strip()

# -- sharepoint (online library) mode --
SP_SITE_URL  = os.getenv("SHAREPOINT_SITE_URL", "https://eychinamanaged.sharepoint.cn/sites/GCSOCR").rstrip("/")
SP_MA_FOLDER = os.getenv("SHAREPOINT_MA_FOLDER", "/sites/GCSOCR/Reporting files templates/1.1 MA整理版")
SP_AR_FOLDER = os.getenv("SHAREPOINT_AR_FOLDER", "/sites/GCSOCR/Reporting files templates/1.2 AR整理版/AR updated Verison")

# -- feishu (Lark) Drive mode --
# Auth is the app-credential model (App ID + App Secret → tenant_access_token), so no
# per-user sign-in is needed: the app reads its own/shared Drive folders. Upload the
# .docx into two Drive folders, share each folder with the app as a reader, and put
# the folder tokens below. FEISHU_APP_SECRET is a secret — keep it out of any
# distributed plaintext (bake it like the Dify keys before shipping). Use feishu.cn
# (China) or larksuite.com (international) for FEISHU_API_BASE.
FEISHU_API_BASE        = os.getenv("FEISHU_API_BASE", "https://open.feishu.cn/open-apis").rstrip("/")
FEISHU_APP_ID          = os.getenv("FEISHU_APP_ID", "").strip()
FEISHU_APP_SECRET      = os.getenv("FEISHU_APP_SECRET", "").strip()
FEISHU_MA_FOLDER_TOKEN = os.getenv("FEISHU_MA_FOLDER_TOKEN", "").strip()
FEISHU_AR_FOLDER_TOKEN = os.getenv("FEISHU_AR_FOLDER_TOKEN", "").strip()
# Optional: letterheads and template_index.xlsx also change over time, so they can be
# sourced from Feishu too. Each is a folder shared with the app: the letterhead folder
# holds the office letterhead .docx files, the index folder holds template_index.xlsx.
# Leave blank to keep using the bundled copies for these.
FEISHU_LETTERHEAD_FOLDER_TOKEN    = os.getenv("FEISHU_LETTERHEAD_FOLDER_TOKEN", "").strip()
FEISHU_TEMPLATE_INDEX_FOLDER_TOKEN = os.getenv("FEISHU_TEMPLATE_INDEX_FOLDER_TOKEN", "").strip()


def _sp_session():
    """A requests session that reuses the user's existing browser sign-in to
    SharePoint, so no stored credential / app registration is needed. Returns
    (session, diag) where diag is a human-readable note about which/how many cookies
    were loaded — surfaced in the UI warning so a 401 can be diagnosed. Best effort:
    if no cookies are readable the REST call fails auth and we fall back to bundled.

    verify=False is passed explicitly on each request below — setting it only on the
    session is not enough: when the per-request verify is None, requests pulls
    REQUESTS_CA_BUNDLE/CURL_CA_BUNDLE from the environment (set on most corporate
    machines) and that overrides the session value, re-enabling validation against a
    bundle that lacks EY's TLS-inspecting-proxy CA. trust_env stays on so the
    corporate HTTP(S) proxy env vars are still honoured for connectivity."""
    s = requests.Session()
    try:
        import browser_cookie3
    except Exception as e:
        return s, f"browser_cookie3 unavailable ({e})"

    # Match on the registrable domain so we catch the FedAuth/rtFa auth cookies
    # regardless of which *.sharepoint.cn host they were set on.
    host = SP_SITE_URL.split("//", 1)[-1].split("/", 1)[0]
    domain = ".".join(host.split(".")[-2:])  # e.g. sharepoint.cn
    # Try each browser separately and merge whatever we can read. Modern Chrome/Edge
    # on Windows use App-Bound Encryption, which browser_cookie3 frequently cannot
    # decrypt — that shows up here as 0 cookies loaded.
    used, total, names = [], 0, []
    for name in ("edge", "chrome", "chromium", "brave", "firefox"):
        fn = getattr(browser_cookie3, name, None)
        if not fn:
            continue
        try:
            n = 0
            for c in fn(domain_name=domain):
                s.cookies.set_cookie(c)
                names.append(c.name)
                n += 1
            if n:
                used.append(f"{name}={n}")
                total += n
        except Exception:
            continue
    if total:
        # List the cookie names so a 401 can be diagnosed: SharePoint cookie auth
        # needs FedAuth (and usually rtFa). If those aren't present the 6 we read are
        # just tracking/consent cookies and the tenant sees us as signed out.
        has_auth = any(c.lower() in ("fedauth", "rtfa") for c in names)
        diag = (f"loaded {total} {domain} cookie(s) from {', '.join(used)}; "
                f"names=[{', '.join(sorted(set(names)))}]; "
                f"FedAuth/rtFa present={has_auth}")
    else:
        diag = (f"no {domain} cookies readable from any browser — sign in to the "
                "library in your browser, or browser cookie encryption blocked access")
    return s, diag


def _sp_list_docx(session, folder_server_relative):
    """Return [(name, server_relative_url), …] for the .docx in a library folder."""
    esc = folder_server_relative.replace("'", "''")
    url = requests.utils.requote_uri(
        f"{SP_SITE_URL}/_api/web/GetFolderByServerRelativePath(decodedurl='{esc}')"
        f"/Files?$select=Name,ServerRelativeUrl")
    r = session.get(url, headers={"Accept": "application/json;odata=nometadata"},
                    timeout=30, verify=False)
    r.raise_for_status()
    payload = r.json()
    # Tolerate both OData flavours: nometadata → {"value": [...]},
    # verbose → {"d": {"results": [...]}}.
    items = payload.get("value") or payload.get("d", {}).get("results", [])
    out = []
    for f in items:
        name = f.get("Name", "")
        if name.lower().endswith(".docx") and not name.startswith("~$"):
            out.append((name, f.get("ServerRelativeUrl", "")))
    return out


def _sp_download(session, server_relative_url):
    esc = server_relative_url.replace("'", "''")
    url = requests.utils.requote_uri(
        f"{SP_SITE_URL}/_api/web/GetFileByServerRelativePath(decodedurl='{esc}')/$value")
    r = session.get(url, timeout=60, verify=False)
    r.raise_for_status()
    return r.content


@st.cache_resource(show_spinner=False)  # the caller renders a full-page loader instead
def _sync_sharepoint_templates():
    """Download the MA/AR .docx from the SharePoint library into a local cache, once
    per process (memoised by st.cache_resource). Returns
    (ar_dir, ma_dir, source, warning|None); on any failure the dirs point back at the
    bundled templates and a warning string is returned for the UI."""
    cache_root = os.path.join(tempfile.gettempdir(), "soc_report_templates")
    ar_dir = os.path.join(cache_root, "AR")
    ma_dir = os.path.join(cache_root, "MA")
    session, cookie_diag = _sp_session()
    try:
        total = 0
        for folder, dest in ((SP_AR_FOLDER, ar_dir), (SP_MA_FOLDER, ma_dir)):
            files = _sp_list_docx(session, folder)
            os.makedirs(dest, exist_ok=True)
            for name, srurl in files:
                data = _sp_download(session, srurl)
                # A real .docx is a zip ("PK"). Anything else (e.g. an HTML sign-in
                # page from an auth redirect) is rejected so we never feed the
                # template pipeline garbage.
                if not data.startswith(b"PK"):
                    raise ValueError(f"'{name}' did not download as a .docx "
                                     f"({len(data)} bytes; likely a sign-in redirect)")
                with open(os.path.join(dest, name), "wb") as fh:
                    fh.write(data)
                total += 1
        if total == 0:
            raise ValueError("no .docx templates found in the SharePoint folders")
        return (ar_dir, ma_dir, "sharepoint", None)
    except Exception as e:
        return (_BUNDLED_AR_DIR, _BUNDLED_MA_DIR, "bundled-fallback",
                f"Could not fetch templates from SharePoint ({e}) [{cookie_diag}]. "
                "Using bundled templates, which may be out of date.")


# ── Feishu (Lark) Drive template source ─────────────────────────────────────────

def _feishu_token():
    """Exchange the app credentials for a tenant_access_token (the app's identity)."""
    r = requests.post(f"{FEISHU_API_BASE}/auth/v3/tenant_access_token/internal",
                      json={"app_id": FEISHU_APP_ID, "app_secret": FEISHU_APP_SECRET},
                      timeout=30, verify=False)
    r.raise_for_status()
    body = r.json()
    if body.get("code") != 0 or not body.get("tenant_access_token"):
        raise ValueError(f"token request failed: {body.get('code')} {body.get('msg')}")
    return body["tenant_access_token"]


def _feishu_list_files(token, folder_token, exts=(".docx",)):
    """Return (matched, inventory) for a Drive folder, where matched is
    [(name, file_token, modified_time), …] for raw uploaded files (type "file")
    whose name ends in one of *exts*, and inventory is a ["name(type)", …] summary
    of *every* entry seen (for diagnostics). modified_time keys the incremental
    skip-unchanged check in _feishu_fetch_into. Native Feishu docs/sheets are
    skipped — only raw uploads download byte-for-byte."""
    exts = tuple(e.lower() for e in exts)
    headers = {"Authorization": f"Bearer {token}"}
    out, inventory, page_token = [], [], None
    while True:
        params = {"folder_token": folder_token, "page_size": 200}
        if page_token:
            params["page_token"] = page_token
        r = requests.get(f"{FEISHU_API_BASE}/drive/v1/files", headers=headers,
                         params=params, timeout=30, verify=False)
        r.raise_for_status()
        body = r.json()
        if body.get("code") != 0:
            raise ValueError(f"list failed: {body.get('code')} {body.get('msg')}")
        data = body.get("data", {})
        for f in data.get("files", []):
            name = f.get("name", "")
            ftype = f.get("type", "")
            inventory.append(f"{name}({ftype})")
            if (name.lower().endswith(exts) and not name.startswith("~$")
                    and ftype == "file"):
                out.append((name, f.get("token", ""), str(f.get("modified_time", ""))))
        if data.get("has_more") and data.get("next_page_token"):
            page_token = data["next_page_token"]
        else:
            return out, inventory


def _feishu_download(token, file_token):
    headers = {"Authorization": f"Bearer {token}"}
    r = requests.get(f"{FEISHU_API_BASE}/drive/v1/files/{file_token}/download",
                     headers=headers, timeout=60, verify=False)
    r.raise_for_status()
    return r.content


def _feishu_fetch_into(token, folder_token, dest, exts):
    """Sync the raw files matching *exts* from a Feishu folder into *dest*.
    Returns (count, inventory) where count is the number of matched files now
    present. Each payload must be a real Office file (zip → "PK").

    Incremental: a .manifest.json in *dest* records each file's token +
    modified_time, so files already present and unchanged are skipped — only
    new/changed templates download. Files removed at the source are deleted
    locally, so deletions still propagate. What does need downloading runs
    concurrently (each file is its own HTTPS round trip, so the work is
    dominated by request latency, not bytes); a small thread pool collapses
    those serial round trips into a few batches (capped to stay under Feishu's
    API rate limit). The first failure is re-raised so the caller's
    all-or-nothing bundled fallback still applies."""
    files, inventory = _feishu_list_files(token, folder_token, exts)
    os.makedirs(dest, exist_ok=True)
    manifest_path = os.path.join(dest, ".manifest.json")
    try:
        with open(manifest_path) as fh:
            manifest = json.load(fh)
    except (OSError, ValueError):
        manifest = {}

    current = {name: {"token": ft, "mtime": mtime} for name, ft, mtime in files}

    # Drop local files the source no longer has, so deletions propagate.
    for fn in os.listdir(dest):
        if fn != ".manifest.json" and fn not in current:
            os.remove(os.path.join(dest, fn))

    # Only fetch what's missing locally or whose token/modified_time changed.
    stale = [(name, ft) for name, ft, _ in files
             if not os.path.isfile(os.path.join(dest, name))
             or manifest.get(name) != current[name]]

    def _fetch_one(name, file_token):
        data = _feishu_download(token, file_token)
        # A real .docx/.xlsx is a zip ("PK"); reject anything else (e.g. an error
        # JSON) so the pipeline never gets garbage.
        if not data.startswith(b"PK"):
            raise ValueError(f"'{name}' did not download as a valid Office file "
                             f"({len(data)} bytes)")
        with open(os.path.join(dest, name), "wb") as fh:
            fh.write(data)

    if stale:
        with ThreadPoolExecutor(max_workers=min(8, len(stale))) as pool:
            futures = [pool.submit(_fetch_one, name, ft) for name, ft in stale]
            for fut in as_completed(futures):
                fut.result()  # propagate the first download/validation error

    # Persist the manifest only after every download succeeded; on a failure the
    # exception above skips this, so the stale files retry next run.
    with open(manifest_path, "w") as fh:
        json.dump(current, fh)
    return len(files), inventory


@st.cache_resource(show_spinner=False)  # the caller renders a full-page loader instead
def _sync_feishu_templates():
    """Download the MA/AR .docx (and, if configured, the letterheads and
    template_index.xlsx) from Feishu Drive into a local cache, once per process
    (memoised). Returns (ar_dir, ma_dir, source, warning|None, letterhead_dir|None,
    index_path|None). The last two are None unless their folder tokens are set and
    fetched successfully — the caller then keeps the bundled copies. On a MA/AR
    failure everything points back at the bundled templates with a warning."""
    cache_root = os.path.join(tempfile.gettempdir(), "soc_report_templates")
    ar_dir  = os.path.join(cache_root, "AR")
    ma_dir  = os.path.join(cache_root, "MA")
    lh_dir  = os.path.join(cache_root, "letterheads")
    idx_dir = os.path.join(cache_root, "index")
    try:
        if not (FEISHU_APP_ID and FEISHU_APP_SECRET):
            raise ValueError("FEISHU_APP_ID / FEISHU_APP_SECRET not set")
        if not (FEISHU_AR_FOLDER_TOKEN and FEISHU_MA_FOLDER_TOKEN):
            raise ValueError("FEISHU_AR_FOLDER_TOKEN / FEISHU_MA_FOLDER_TOKEN not set")
        token = _feishu_token()
        total, diag = 0, []
        for label, folder_token, dest in (("AR", FEISHU_AR_FOLDER_TOKEN, ar_dir),
                                          ("MA", FEISHU_MA_FOLDER_TOKEN, ma_dir)):
            n, inventory = _feishu_fetch_into(token, folder_token, dest, (".docx",))
            total += n
            diag.append(f"{label}: {len(inventory)} item(s) "
                        + ("[" + ", ".join(inventory[:10]) + "]" if inventory else "[empty]"))
        if total == 0:
            raise ValueError("no uploaded .docx files found — Feishu may have converted "
                             "your Word files to native docs (need type 'file', not 'docx'/'doc'), "
                             "or the folder token is wrong / not shared with the app. "
                             "Folder contents — " + "; ".join(diag))

        # Optional extras. Best effort: a failure here keeps the MA/AR result and
        # falls back to the bundled copy for just that asset, noted in the warning.
        lh_ret, idx_ret, extra = None, None, []
        if FEISHU_LETTERHEAD_FOLDER_TOKEN:
            try:
                n, inv = _feishu_fetch_into(token, FEISHU_LETTERHEAD_FOLDER_TOKEN, lh_dir, (".docx",))
                if n:
                    lh_ret = lh_dir
                else:
                    extra.append("letterheads: no uploaded .docx found ("
                                 + (", ".join(inv[:10]) if inv else "empty") + ")")
            except Exception as e:
                extra.append(f"letterheads: {e}")
        if FEISHU_TEMPLATE_INDEX_FOLDER_TOKEN:
            try:
                n, inv = _feishu_fetch_into(token, FEISHU_TEMPLATE_INDEX_FOLDER_TOKEN, idx_dir, (".xlsx",))
                if n:
                    xs = [f for f in os.listdir(idx_dir) if f.lower().endswith(".xlsx")]
                    pick = "template_index.xlsx" if "template_index.xlsx" in xs else xs[0]
                    idx_ret = os.path.join(idx_dir, pick)
                else:
                    extra.append("template_index: no uploaded .xlsx found ("
                                 + (", ".join(inv[:10]) if inv else "empty") + ")")
            except Exception as e:
                extra.append(f"template_index: {e}")

        warning = ("Loaded MA/AR from Feishu, but " + "; ".join(extra)
                   + " — using bundled copies for those.") if extra else None
        return (ar_dir, ma_dir, "feishu", warning, lh_ret, idx_ret)
    except Exception as e:
        return (_BUNDLED_AR_DIR, _BUNDLED_MA_DIR, "bundled-fallback",
                f"Could not fetch templates from Feishu ({e}). Using bundled "
                "templates, which may be out of date.", None, None)


def _resolve_template_dirs():
    """Decide where MA/AR templates are read from this session.

    Returns (ar_dir, ma_dir, source, warning|None, letterhead_dir|None,
    index_path|None). The last two are non-None only when Feishu supplies them;
    otherwise the caller keeps the bundled letterheads / template_index.xlsx.
    Honours TEMPLATE_SOURCE ('sharepoint' / 'feishu' / 'onedrive' / 'bundled'); any
    miss falls back to the bundled templates with a warning string for the UI. NOTE:
    call this only after st.set_page_config — the online paths may render a spinner.
    """
    if TEMPLATE_SOURCE == "feishu":
        return _sync_feishu_templates()
    if TEMPLATE_SOURCE == "sharepoint":
        ar, ma, src, warn = _sync_sharepoint_templates()
        return (ar, ma, src, warn, None, None)
    if TEMPLATE_SOURCE == "onedrive" and TEMPLATE_BASE_PATH:
        ar = os.path.join(TEMPLATE_BASE_PATH, *AR_TEMPLATE_SUBPATH.replace("\\", "/").split("/"))
        ma = os.path.join(TEMPLATE_BASE_PATH, *MA_TEMPLATE_SUBPATH.replace("\\", "/").split("/"))
        missing = [p for p in (ar, ma) if not os.path.isdir(p)]
        if not missing:
            return (ar, ma, "onedrive", None, None, None)
        return (_BUNDLED_AR_DIR, _BUNDLED_MA_DIR, "bundled-fallback",
                "Synced template folder not found — using bundled templates, which "
                "may be out of date. Missing: " + "; ".join(missing), None, None)
    return (_BUNDLED_AR_DIR, _BUNDLED_MA_DIR, "bundled", None, None, None)


# EY office letterhead .docx files (downloaded from the EY Templates Word add-in,
# grouped + centred in Word). The header on the Auditor's Report pages is taken
# from whichever one the user picks. Listed at runtime, so end users can add or
# remove office letterheads next to the .exe without a rebuild. This is the bundled
# default; TEMPLATE_SOURCE=feishu can override it below if a letterhead folder token
# is configured.
LETTERHEAD_DIR  = os.path.join(_BUNDLED_AR_DIR, "letterheads")
EY_FIRM_NAME    = "Ernst & Young Hua Ming LLP"


def list_letterheads():
    """Return the sorted list of letterhead .docx filenames in LETTERHEAD_DIR.

    Read fresh on every call so the dropdown reflects whatever files currently
    sit in the folder (alongside the .exe when frozen)."""
    try:
        return sorted(
            f for f in os.listdir(LETTERHEAD_DIR)
            if f.lower().endswith(".docx") and not f.startswith("~$")
        )
    except (FileNotFoundError, NotADirectoryError):
        return []

API_BASE_URL  = os.getenv("DIFY_API_BASE_URL", "https://api.dify.ai/v1")
API_KEY_MAIN  = os.getenv("DIFY_API_KEY_MAIN", "")
API_KEY_SUB1  = os.getenv("DIFY_API_KEY_SUB1", "")
API_KEY_SUB2  = os.getenv("DIFY_API_KEY_SUB2", "")

st.set_page_config(page_title="AI-Driven Report Generation", layout="wide")

# Resolve where MA/AR templates come from this session (bundled / SharePoint /
# Feishu / synced folder). resolve_template() and the UI read the module globals
# set below. Feishu mode may also override the letterhead dir and
# template_index.xlsx path; any value it leaves None keeps the bundled copy.
#
# The online sources (Feishu / SharePoint) download at startup, which can take a
# few seconds. The fetch blocks the script, and Streamlit keeps the *previous*
# render on screen (dimmed) while it blocks — so a plain spinner ends up sitting on
# top of the old form (noisy, see appear.png). Instead we paint a full-viewport
# overlay just before the blocking call: it covers everything underneath, then we
# remove it once templates are ready and let the page render. Stashed in
# session_state so the overlay shows only on the first resolve of the session;
# bundled / onedrive resolve instantly and never show it.
if TEMPLATE_SOURCE in ("feishu", "sharepoint") and "_template_dirs" not in st.session_state:
    _src_label = "SharePoint" if TEMPLATE_SOURCE == "sharepoint" else "Feishu"
    _overlay = st.empty()
    _overlay.markdown(
        f"""
        <style>@keyframes soc-spin {{ to {{ transform: rotate(360deg); }} }}</style>
        <div style="position:fixed; inset:0; z-index:2147483647;
                    display:flex; flex-direction:column; gap:1.1rem;
                    align-items:center; justify-content:center; text-align:center;
                    background:var(--background-color, #0e1117);
                    color:var(--text-color, #fafafa);">
          <div style="width:46px; height:46px; border-radius:50%;
                      border:4px solid rgba(128,128,128,.35); border-top-color:#ff4b4b;
                      animation:soc-spin 1s linear infinite;"></div>
          <div style="font-size:1.15rem; font-weight:600;">
            Fetching the latest templates from {_src_label}…</div>
          <div style="opacity:.7;">The report form will appear once they're ready.</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.session_state["_template_dirs"] = _resolve_template_dirs()
    _overlay.empty()  # drop the overlay; the page renders below in this same run

st.title("AI-Driven SOC Report Generation")

(AR_TEMPLATE_DIR, MA_TEMPLATE_DIR, TEMPLATE_DIR_SOURCE, TEMPLATE_DIR_WARNING,
 _FE_LETTERHEAD_DIR, _FE_INDEX_PATH) = st.session_state.get(
    "_template_dirs") or _resolve_template_dirs()
if _FE_LETTERHEAD_DIR:
    LETTERHEAD_DIR = _FE_LETTERHEAD_DIR
if _FE_INDEX_PATH:
    TEMPLATE_INDEX = _FE_INDEX_PATH

# ── API config — loaded from bundled .env, never shown in UI ──────────────────
api_base = API_BASE_URL
key_main = API_KEY_MAIN
key_sub1 = API_KEY_SUB1
key_sub2 = API_KEY_SUB2

# ── Sidebar ────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.header("SOC Report Generator")
    st.markdown("AI-Driven Report Generation")
    st.markdown("---")
    if st.button("🔄 Reset All Steps", use_container_width=True):
        # Clear the generated results and input snapshots. Once they're gone the
        # form re-renders empty (it lives under `if not final_done`, so its widget
        # state was discarded while the report was on screen) and the user starts
        # fresh.
        _was_complete = "final_result" in st.session_state
        for k in ["main_outputs", "sub1_outputs", "final_result", "user_inputs",
                  "template_config", "ma_ar_only", "final_bytes", "final_filename"]:
            st.session_state.pop(k, None)
        # CUEC/UER keep a shadow `pref_…` copy so the selection survives a Reset
        # done mid-run or after a MA+AR-only file (the form is still on screen, like
        # every other field). After a COMPLETE report the form was hidden and all
        # other fields reset, so here — and only here — drop the CUEC/UER shadows
        # and widget keys too, so they default along with everything else.
        if _was_complete:
            for k in [k for k in st.session_state
                      if "form_is_cuec" in k or "form_is_uer" in k]:
                st.session_state.pop(k, None)
        st.rerun()

    # The online fetch is memoised for the life of the process, so a long-running
    # session won't pick up templates updated at the source mid-run. This button
    # drops the cache and re-downloads on the next render.
    if TEMPLATE_SOURCE in ("sharepoint", "feishu"):
        _src_label = "SharePoint" if TEMPLATE_SOURCE == "sharepoint" else "Feishu"
        if st.button(f"⬇️ Refresh templates from {_src_label}", use_container_width=True):
            (_sync_sharepoint_templates if TEMPLATE_SOURCE == "sharepoint"
             else _sync_feishu_templates).clear()
            st.session_state.pop("_template_dirs", None)  # re-show the fetch loader
            st.rerun()

# ── Progress indicator ─────────────────────────────────────────────────────────
main_done  = "main_outputs"  in st.session_state
sub1_done  = "sub1_outputs"  in st.session_state
final_done = "final_result"  in st.session_state

s1 = "✅" if main_done  else "🔵"
s2 = "✅" if sub1_done  else ("🟡" if main_done  else "⚪")
s3 = "✅" if final_done else ("🟡" if sub1_done  else "⚪")

def _status_html(a, b, c):
    return (
        "<div style='display:flex;gap:2rem;padding:0.5rem 0 1.2rem 0;font-size:1rem'>"
        f"<span>{a} <b>Step 1</b> — MAIN: Extract &amp; Prepare</span>"
        "<span>→</span>"
        f"<span>{b} <b>Step 2</b> — SUB1: Entity Level</span>"
        "<span>→</span>"
        f"<span>{c} <b>Step 3</b> — SUB2: Final Report</span>"
        "</div>"
    )

status_bar = st.empty()
status_bar.markdown(_status_html(s1, s2, s3), unsafe_allow_html=True)

# ── Template helpers ───────────────────────────────────────────────────────────

def get_standard_options(report_type):
    """Return the list of applicable attestation standards for a given report type."""
    if report_type.startswith("SOC1"):
        return ["SSAE 18", "ISAE 3402", "SSAE 18 & ISAE 3402 Combined"]
    else:
        return ["SSAE 18", "ISAE 3000", "SSAE 18 & ISAE 3000 Combined"]


def resolve_template(report_type, standard, sso, language, sheet):
    """
    Look up template_index.xlsx for a matching row and return (wp_no, filepath|None).
    On error returns (None, error_message_string) so callers can surface the problem.
    sheet must be 'AR' or 'MA'.
    """
    # Map UI values → spreadsheet values
    if report_type.startswith("SOC1"):
        category = "SOC 1"
    else:
        category = "SOC 2"

    if "TYPE1" in report_type:
        typ = "Type I"
    else:
        typ = "Type II"

    if "Combined" in standard:
        std_mapped = "Combined"
    else:
        std_mapped = standard  # "SSAE 18", "ISAE 3402", "ISAE 3000"

    sso_map = {"None": "none", "All carve out": "all carve out", "Inclusive": "Inclusive"}
    sso_mapped = sso_map.get(sso, "none")

    lang_map = {"English": "EN", "中文": "CN"}
    lang_mapped = lang_map.get(language, "EN")

    template_dir = AR_TEMPLATE_DIR if sheet == "AR" else MA_TEMPLATE_DIR

    try:
        # Do NOT use read_only=True — it can silently fail to iterate rows in
        # some environments.
        wb = openpyxl.load_workbook(TEMPLATE_INDEX, data_only=True)
        ws = wb[sheet]
        rows = list(ws.iter_rows(values_only=True))
        wb.close()
    except Exception as e:
        return (None, f"Cannot load template index: {e}")

    if not rows:
        return (None, "Template index sheet is empty")

    header = list(rows[0])
    try:
        col_cat  = header.index("Category")
        col_type = header.index("Type")
        col_std  = header.index("Standards")
        col_sso  = header.index("Sub-service Organization (SSO)")
        col_lang = header.index("Language")
        col_wp   = next(
            i for i, h in enumerate(header)
            if h and str(h).strip().upper().startswith("WP")
        )
    except (ValueError, StopIteration) as e:
        return (None, f"Template index column not found: {e}")

    for row in rows[1:]:
        if (str(row[col_cat]  or "").strip() == category  and
                str(row[col_type] or "").strip() == typ       and
                str(row[col_std]  or "").strip() == std_mapped and
                str(row[col_sso]  or "").strip() == sso_mapped and
                str(row[col_lang] or "").strip() == lang_mapped):
            wp_val = row[col_wp]
            if wp_val is None:
                # This combination has no template — keep iterating in case
                # another row matches (shouldn't happen, but safe).
                continue
            wp_no = str(wp_val).strip()
            try:
                for fname in sorted(os.listdir(template_dir)):
                    if fname.endswith(".docx") and fname.startswith(wp_no + " "):
                        return (wp_no, os.path.join(template_dir, fname))
            except OSError as e:
                return (wp_no, f"Cannot list template directory: {e}")
            return (wp_no, None)

    return (None, None)


def _normalize_ws(s):
    """Collapse all whitespace sequences to a single space and strip."""
    return " ".join(s.split())


_MONTHS = [
    "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December",
]


def _format_date(s, language="English"):
    """Convert YYYY-MM-DD (also YYYY/M/D, YYYY.M.D, 1-digit month/day) to
    'Month D, YYYY' (EN) or 'YYYY年M月D日' (CN).
    Strings that do not match a recognised date pattern are returned unchanged."""
    if not s:
        return s
    m = re.match(r"^(\d{4})[-/.](\d{1,2})[-/.](\d{1,2})$", s.strip())
    if m:
        y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if 1 <= mo <= 12 and 1 <= d <= 31:
            if language == "中文":
                return f"{y}年{mo}月{d}日"
            return f"{_MONTHS[mo - 1]} {d}, {y}"
    return s


def _capitalize_name(s):
    """Normalise a typed name to word-initial capitals.

    - ALL UPPERCASE input is title-cased ('ACME TECH CO., LTD.'
      → 'Acme Tech Co., Ltd.').
    - Otherwise each word starting with a lowercase letter is capitalised,
      keeping the rest of the word as typed ('fgde information Technology'
      → 'Fgde Information Technology'). Words with internal capitals
      ('iPhone') and words already capitalised are left unchanged.
    - Names containing CJK characters are never changed."""
    if not s:
        return s
    letters = [c for c in s if c.isalpha()]
    if not letters or not all(c.isascii() for c in letters):
        return s
    if all(c.isupper() for c in letters):
        return re.sub(
            r"[A-Za-z][A-Za-z'’]*",
            lambda m: m.group(0)[0].upper() + m.group(0)[1:].lower(),
            s,
        )

    def _fix_word(m):
        w = m.group(0)
        if w[0].islower() and not any(c.isupper() for c in w[1:]):
            return w[0].upper() + w[1:]
        return w

    return re.sub(r"[A-Za-z][A-Za-z'’]*", _fix_word, s)


def _kw_in(comment_text, keyword):
    """Check if keyword appears in comment_text, ignoring all internal whitespace.
    Stripping spaces handles XML-inserted spaces between CJK characters."""
    return keyword.replace(" ", "") in comment_text.replace(" ", "")


# Comment keyword sets for deletion logic
_CUEC_KW     = ["无用户补充", "未识别用户补充", "user entity补充", "无用户实体补充"]
_SSO_CC_KW   = ["子服务机构补偿", "子服务机构补充"]
_TRANS_KW    = ["处理transaction", "processing user entity transaction"]
_SINGLE_KW   = ["single user entity report", "single user entity时", "single user entity 时"]
_OTHER_KW    = ["other information"]  # kept unless user says no Other Information section
_AI_KW       = ["使用到了AI技术", "subject matter中某部分使用到了"]  # AI scope exclusion paragraph


def _comment_span_text(para_el, cid, ns_w, id_attr):
    """Return the run text covered by comment *cid* inside one paragraph
    element, i.e. the text between its commentRangeStart/End markers.
    Deleted text (w:delText) is naturally excluded and inserted text
    (w:t inside w:ins) included, matching accept-tracked-changes cleaning."""
    inside = False
    parts = []
    for el in para_el.iter():
        if el.tag == f"{{{ns_w}}}commentRangeStart" and el.get(id_attr) == cid:
            inside = True
        elif el.tag == f"{{{ns_w}}}commentRangeEnd" and el.get(id_attr) == cid:
            inside = False
        elif inside and el.tag == f"{{{ns_w}}}t":
            parts.append(el.text or "")
    return "".join(parts)


def _build_annotation_maps(docx_bytes, flags):
    """
    Parse the docx XML (from raw bytes) to determine per-paragraph actions.

    Returns:
        del_indices     — set of body-child indices to remove entirely
        single_ue_indices — set of body-child indices where [..] bracket
                            content is conditionally deleted (single user entity)
        span_del_texts  — list of (body-child index, text) pairs: the exact
                          commented sentence span to remove from inside a
                          paragraph that otherwise stays (single-UE comments
                          anchored to one sentence of a longer paragraph)
    """
    ns_w    = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    p_tag   = f"{{{ns_w}}}p"
    id_attr = f"{{{ns_w}}}id"

    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as z:
            doc_xml = z.read("word/document.xml").decode("utf-8")
            if "word/comments.xml" in z.namelist():
                comments_xml = z.read("word/comments.xml").decode("utf-8")
            else:
                return set(), set(), []
    except Exception:
        return set(), set(), []

    if doc_xml.startswith("\ufeff"):
        doc_xml = doc_xml[1:]
    if comments_xml.startswith("\ufeff"):
        comments_xml = comments_xml[1:]

    # Build comment id → normalised lower-case text map
    try:
        root_c = ET.fromstring(comments_xml)
    except ET.ParseError:
        return set(), set(), []
    comment_texts = {}
    for comment in root_c.findall(f"{{{ns_w}}}comment"):
        cid = comment.get(id_attr)
        texts = [t.text or "" for t in comment.findall(f".//{{{ns_w}}}t")]
        comment_texts[cid] = _normalize_ws(" ".join(texts)).lower()

    # Build comment id → set of body-child indices (paragraphs only)
    try:
        root_d = ET.fromstring(doc_xml)
    except ET.ParseError:
        return set(), set(), []
    body = root_d.find(f"{{{ns_w}}}body")
    if body is None:
        return set(), set(), []

    comment_to_indices = {}
    for i, child in enumerate(body):
        if child.tag != p_tag:
            continue
        for crs in child.findall(f".//{{{ns_w}}}commentRangeStart"):
            cid = crs.get(id_attr)
            if cid:
                comment_to_indices.setdefault(cid, set()).add(i)

    del_indices      = set()
    single_ue_indices = set()
    span_del_texts   = []

    body_list = list(body)

    for cid, text in comment_texts.items():
        # "Other Information" paragraphs (注：若无Other Information这一章节，则删除…):
        # kept by default; deleted only when the user indicated the report has
        # no Other Information section and the comment is a delete instruction.
        if any(_kw_in(text, kw.lower()) for kw in _OTHER_KW):
            if (not flags.get("has_other_information", True)
                    and ("删除" in text or "delete" in text.lower())):
                del_indices |= comment_to_indices.get(cid, set())
            continue

        indices = comment_to_indices.get(cid, set())
        if not indices:
            continue

        # Single-UE: if comment says "delete" → when the flag is set, delete
        # only the commented sentence span if the comment covers part of a
        # longer paragraph, or the whole paragraph when it covers all of it
        # (or explicitly says 本段). Otherwise → bracket-replacement path.
        if any(_kw_in(text, kw.lower()) for kw in _SINGLE_KW):
            is_delete_cmd = "删除" in text or "delete" in text.lower()
            if is_delete_cmd and flags.get("single_user_entity", False):
                for idx in indices:
                    child = body_list[idx]
                    span = _comment_span_text(child, cid, ns_w, id_attr)
                    span_n = _normalize_ws(span)
                    para_n = _normalize_ws("".join(
                        t.text or "" for t in child.findall(f".//{{{ns_w}}}t")
                    ))
                    # Partial-paragraph sentence deletion only when the span is
                    # a substantial fragment (≥30 chars guards against ranges
                    # degenerately anchored to a single word like "The"/"描述")
                    # and the comment does not explicitly target the paragraph.
                    if (span_n and span_n != para_n and len(span_n) >= 30
                            and not _kw_in(text, "本段")):
                        span_del_texts.append((idx, span))
                    else:
                        del_indices.add(idx)
            else:
                single_ue_indices |= indices
            continue

        # AI scope-exclusion paragraph
        if any(_kw_in(text, kw.lower()) for kw in _AI_KW):
            if not flags.get("has_ai_scope_exclusion", False):
                del_indices |= indices
            continue

        should_delete = False
        if not flags.get("cuec_identified", True):
            if any(_kw_in(text, kw.lower()) for kw in _CUEC_KW):
                should_delete = True
        if not flags.get("sso_cc_identified", True):
            if any(_kw_in(text, kw.lower()) for kw in _SSO_CC_KW):
                should_delete = True
        if not flags.get("has_transaction_processing", True):
            if any(_kw_in(text, kw.lower()) for kw in _TRANS_KW):
                # Only delete SHORT standalone paragraphs (≤200 chars after
                # stripping whitespace). Long paragraphs containing the
                # transaction phrase as just one embedded clause are handled
                # by inline text substitution in build_substitutions instead.
                for idx in indices:
                    if idx < len(body_list) and body_list[idx].tag == p_tag:
                        para_txt = "".join(
                            t.text or ""
                            for t in body_list[idx].findall(f".//{{{ns_w}}}t")
                        ).strip()
                        if len(para_txt) <= 200:
                            del_indices.add(idx)
                # (do not set should_delete — individual indices already handled)

        if should_delete:
            del_indices |= indices

    return del_indices, single_ue_indices, span_del_texts


def _reject_format_changes(xml_str):
    """
    Reject tracked formatting changes by restoring the OLD formatting:
    - w:pPrChange: replace the parent w:pPr with the OLD w:pPr inside pPrChange
    - w:rPrChange: replace the parent w:rPr with the OLD w:rPr inside rPrChange

    In EY templates, pPrChange/rPrChange record authoring edits that should
    NOT be accepted into the final output — rejecting restores the intended
    original formatting (e.g., list style a. instead of Wingdings bullet l,
    non-bold run instead of bold run).
    """
    def _reject_one(xml_s, change_tag, parent_tag):
        change_re = re.compile(
            rf"<{re.escape(change_tag)}\b[^>]*>(.*?)</{re.escape(change_tag)}>",
            re.DOTALL,
        )
        result = xml_s
        for m in reversed(list(change_re.finditer(result))):
            chg_start, chg_end = m.start(), m.end()

            # Extract old parent element from inside the change block
            old_m = re.search(
                rf"<{re.escape(parent_tag)}\b[^>]*>.*?</{re.escape(parent_tag)}>",
                m.group(1), re.DOTALL,
            )

            # The outer closing tag immediately follows the change block
            rest = result[chg_end:]
            close_m = re.match(rf"\s*</{re.escape(parent_tag)}>", rest)
            if not close_m:
                # Unexpected structure — just drop the change block
                result = result[:chg_start] + result[chg_end:]
                continue
            outer_end = chg_end + close_m.end()

            # The LAST opening parent tag before the change block
            before = result[:chg_start]
            opens = list(re.finditer(rf"<{re.escape(parent_tag)}\b[^>]*>", before))
            if not opens:
                result = result[:chg_start] + result[chg_end:]
                continue
            outer_start = opens[-1].start()

            replacement = old_m.group(0) if old_m else ""
            result = result[:outer_start] + replacement + result[outer_end:]

        return result

    xml_str = _reject_one(xml_str, "w:pPrChange", "w:pPr")
    xml_str = _reject_one(xml_str, "w:rPrChange", "w:rPr")
    return xml_str


def _apply_xml_cleaning(xml_str):
    """
    Apply all tracked-change cleaning to a raw XML string.

    Strategy (matches EY template authoring conventions):
      - w:ins  → ACCEPT:  unwrap, keep content
      - w:del  → ACCEPT:  remove deleted content entirely
      - pPrChange / rPrChange → REJECT:  restore OLD formatting
      - Table / section format-change markers → removed
      - Comment markers → removed
    """
    xml_str = xml_str.lstrip("\ufeff")

    # 1. Accept deletions: remove <w:del>…</w:del> blocks entirely.
    # (?<!/) ensures self-closing <w:del/> markers are NOT matched here
    # (they are handled by step 2); without this guard the regex consumes
    # from <w:del/> all the way to the next unrelated </w:del>, deleting
    # large valid content spans.
    xml_str = re.sub(
        r"<w:del\b[^>]*(?<!/)>.*?</w:del>", "",
        xml_str, flags=re.DOTALL,
    )

    # 2. Remove self-closing tracked-change markers
    xml_str = re.sub(r"<w:del\b[^>]*/>",  "", xml_str)
    xml_str = re.sub(r"<w:ins\b[^>]*/>",  "", xml_str)

    # 3. Accept insertions: unwrap <w:ins>…</w:ins>
    xml_str = re.sub(
        r"<w:ins\b[^>]*(?<!/)>(.*?)</w:ins>", r"\1",
        xml_str, flags=re.DOTALL,
    )

    # 5. Reject format changes: restore OLD pPr / rPr
    xml_str = _reject_format_changes(xml_str)

    # 6. Remove table / section format-change tracking elements
    for _ftag in ("w:tblPrChange", "w:trPrChange", "w:tcPrChange", "w:sectPrChange",
                  "w:numChange"):
        xml_str = re.sub(
            rf"<{_ftag}\b[^>]*>.*?</{_ftag}>", "",
            xml_str, flags=re.DOTALL,
        )

    # 7. Remove comment range markers
    xml_str = re.sub(r"<w:commentRangeStart[^>]*/>", "", xml_str)
    xml_str = re.sub(r"<w:commentRangeEnd[^>]*/>",   "", xml_str)

    # 8. Remove runs whose only non-rPr child is a comment reference
    xml_str = re.sub(
        r"<w:r\b[^>]*>(?:(?!</w:r>).)*?<w:commentReference[^>]*/>\s*</w:r>",
        "", xml_str, flags=re.DOTALL,
    )

    return xml_str


def _clean_docx_bytes(docx_bytes):
    """
    Accept/reject tracked changes and clear comments in a docx file.

    Uses _apply_xml_cleaning() on word/document.xml and related XML parts
    (headers, footers, footnotes, endnotes).  word/comments.xml is replaced
    with an empty comments document so Word opens without comment balloons.
    """
    _EMPTY_COMMENTS = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:comments xmlns:w="http://schemas.openxmlformats.org/'
        'wordprocessingml/2006/main"></w:comments>'
    ).encode("utf-8")

    buf_in  = io.BytesIO(docx_bytes)
    buf_out = io.BytesIO()

    with zipfile.ZipFile(buf_in, "r") as zin:
        with zipfile.ZipFile(buf_out, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data  = zin.read(info.filename)
                fname = info.filename
                if fname == "word/comments.xml":
                    data = _EMPTY_COMMENTS
                elif (
                    fname == "word/document.xml"
                    or re.match(r"word/(header|footer)\d+\.xml$", fname)
                    or fname in ("word/footnotes.xml", "word/endnotes.xml")
                ):
                    data = _apply_xml_cleaning(data.decode("utf-8")).encode("utf-8")
                zout.writestr(info, data)

    buf_out.seek(0)
    return buf_out.read()


def _smart_replace_in_para(para, old, new):
    """Replace one occurrence of *old* with *new* in the paragraph, modifying
    only the minimal run span that contains the match.  Per-run formatting
    (italic, bold, size, etc.) on runs outside that span is fully preserved.
    Returns True if a replacement was made, False otherwise."""
    if not old or not para.runs:
        return False
    texts = [r.text for r in para.runs]
    full  = "".join(texts)
    if old not in full:
        return False
    pos     = full.index(old)
    end_pos = pos + len(old)
    # Cumulative start offset of each run within full
    cum = 0
    starts = []
    for t in texts:
        starts.append(cum)
        cum += len(t)
    # First run whose text overlaps the match
    fi = next((i for i in range(len(texts)) if starts[i] + len(texts[i]) > pos), None)
    if fi is None:
        return False
    # Last run whose text overlaps the match
    li = next((i for i in range(len(texts) - 1, -1, -1) if starts[i] < end_pos), None)
    if li is None:
        return False
    prefix = texts[fi][: pos - starts[fi]]
    suffix = texts[li][end_pos - starts[li] :]
    para.runs[fi].text = prefix + new + suffix
    for k in range(fi + 1, li + 1):
        para.runs[k].text = ""
    return True


def _set_outline_level(para, level):
    """Tag a paragraph with an explicit w:outlineLvl (0 = 1st level) so it
    shows in Word's navigation pane. Heading-style references alone do not
    survive the merge into the EY template document (the styles are not
    defined there), and outlineLvl has no effect on visual formatting."""
    pPr = para._p.get_or_add_pPr()
    ol = pPr.find(qn("w:outlineLvl"))
    if ol is None:
        ol = OxmlElement("w:outlineLvl")
        tail = next(
            (ch for ch in pPr
             if ch.tag in (qn("w:rPr"), qn("w:sectPr"), qn("w:pPrChange"))),
            None,
        )
        if tail is not None:
            tail.addprevious(ol)
        else:
            pPr.append(ol)
    ol.set(qn("w:val"), str(level))


def fill_and_process_template(template_path, subs, flags, language="English"):
    """
    Process an EY MA/AR docx template:
      1. Compute comment-based annotation maps from the original file.
      2. Accept tracked changes and clear comments (regex on raw XML).
      3. Apply placeholder substitutions.
      4. Handle inline square brackets:
           - single-user-entity annotated paras: delete [..] when flag set,
             otherwise strip the brackets and keep the content.
           - [or ..] alternative phrases: always removed.
           - remaining [..] brackets: brackets stripped, content kept.
      5. Delete wholly-conditional paragraphs (CUEC / SSO CC / transaction /
         AI-scope when not applicable).
      6. Standardise fonts: Times New Roman (EN) or 华文楷体 (CN), 11 pt,
         bold removed, italic preserved.

    Returns the modified document as bytes.
    """
    # ── Step 1: annotation maps from the original (comments still present) ──
    with open(template_path, "rb") as fh:
        raw_bytes = fh.read()
    del_indices, single_ue_indices, span_del_texts = _build_annotation_maps(
        raw_bytes, flags
    )

    # ── Step 2: accept track changes + clear comments ──────────────────────
    cleaned_bytes = _clean_docx_bytes(raw_bytes)
    doc = Document(io.BytesIO(cleaned_bytes))

    # ── Step 2b: delete single-UE commented sentence spans ─────────────────
    # Runs BEFORE placeholder substitution because the extracted span text is
    # the raw template wording, which may still contain placeholders.
    if span_del_texts:
        _body_kids = list(doc.element.body)
        _para_by_el = {p._element: p for p in doc.paragraphs}
        for _idx, _span in span_del_texts:
            if _idx >= len(_body_kids):
                continue
            _para = _para_by_el.get(_body_kids[_idx])
            if _para is None:
                continue
            if not _smart_replace_in_para(_para, _span, ""):
                _smart_replace_in_para(_para, _span.strip(), "")
            # Drop any leading whitespace left when the deleted sentence was
            # at the start of the paragraph.
            for _r in _para.runs:
                if _r.text:
                    _r.text = _r.text.lstrip()
                    break

    # ── Step 3: placeholder substitutions ──────────────────────────────────
    def _apply_subs_to_para(para):
        if not para.runs:
            return
        # Phase 0: compound date-range placeholders contain the single 【日期】/
        # [date] placeholder — consume them first with cross-run replacement,
        # otherwise phase 1 could fill one half with the single date when the
        # range spans multiple runs.
        for placeholder in ("[date] to [date]", "【日期】至【日期】"):
            value = subs.get(placeholder)
            if value is not None:
                while _smart_replace_in_para(para, placeholder, value):
                    pass
        # Phase 1: per-run substitution — preserves individual run formatting
        # (bold, italic, etc.) when the placeholder is entirely within one run.
        for run in para.runs:
            for placeholder, value in subs.items():
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)
        # Phase 2: smart span replacement for cross-run placeholders.
        # Loop each placeholder until no occurrences remain — a single paragraph
        # may contain the same placeholder more than once (e.g. [Service
        # organization short name] appears multiple times in the MA description).
        for placeholder, value in subs.items():
            while _smart_replace_in_para(para, placeholder, value):
                pass
        # Phase 3: normalize consecutive spaces that may arise from empty
        # substitutions, both within a run and across run boundaries.
        prev_ended_space = False
        for run in para.runs:
            if run.text:
                run.text = re.sub(r"  +", " ", run.text)
                if prev_ended_space and run.text.startswith(" "):
                    run.text = run.text.lstrip(" ")
                prev_ended_space = run.text.endswith(" ")
            # empty run: prev_ended_space unchanged (gap is invisible)

    for para in doc.paragraphs:
        _apply_subs_to_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _apply_subs_to_para(para)

    # ── Step 4: inline bracket handling ────────────────────────────────────
    body_children  = list(doc.element.body)
    body_child_map = {child: i for i, child in enumerate(body_children)}
    _p_tag         = qn("w:p")
    _single_ue_flag = flags.get("single_user_entity", False)

    def _apply_brackets(para, is_single_ue):
        # Process one bracket match per loop iteration using smart replacement so
        # that per-run italic/bold formatting outside the matched span is preserved.
        changed = True
        while changed:
            changed = False
            full_text = "".join(r.text for r in para.runs)
            if not full_text or ("[" not in full_text and "【" not in full_text):
                break

            # a) [or …] / 【或…】 alternative phrases — always remove
            #    (with any leading space for the EN variant)
            m = re.search(r" ?\[or [^\]]+\]", full_text, flags=re.IGNORECASE)
            if m:
                _smart_replace_in_para(para, m.group(0), "")
                changed = True
                continue
            m = re.search(r"【或[^】]*】", full_text)
            if m:
                _smart_replace_in_para(para, m.group(0), "")
                changed = True
                continue

            # b) single-user-entity annotated paragraphs
            if is_single_ue:
                if _single_ue_flag:
                    # delete bracketed content entirely
                    m = re.search(r"\[[^\]]*\]|【[^】]*】", full_text)
                    if m:
                        _smart_replace_in_para(para, m.group(0), "")
                        changed = True
                        continue
                else:
                    # strip brackets, keep content
                    m = re.search(r"\[([^\]]*)\]|【([^】]*)】", full_text)
                    if m:
                        _smart_replace_in_para(para, m.group(0), m.group(1) or m.group(2) or "")
                        changed = True
                        continue

            # c) any remaining [..] / 【..】 — strip brackets, keep content
            m = re.search(r"\[([^\]]*)\]|【([^】]*)】", full_text)
            if m:
                _smart_replace_in_para(para, m.group(0), m.group(1) or m.group(2) or "")
                changed = True
                continue

        # d) Orphaned brackets from template-authoring inconsistencies. Some
        #    multi-SSO templates open the conditional block by merging its 【 into
        #    the first 【服务机构简称】 placeholder (single bracket instead of
        #    double); once that placeholder is substituted away, the block's
        #    closing 】 is left with no opening to pair against, so cases a–c above
        #    can never match it. By this point all real bracket pairs and [or…]
        #    phrases are gone, so any remaining 【】[] is an artifact — strip the
        #    lone bracket characters (these templates never use them as literal
        #    text). See template "12.2 AR_SOC2 Type I_SSAE18_IL503_CN".
        for run in para.runs:
            if run.text and re.search(r"[【】\[\]]", run.text):
                run.text = re.sub(r"[【】\[\]]", "", run.text)

        # Normalize spaces introduced by empty removals
        prev_ended_space = False
        for run in para.runs:
            if run.text:
                run.text = re.sub(r"  +", " ", run.text)
                if prev_ended_space and run.text.startswith(" "):
                    run.text = run.text.lstrip(" ")
                prev_ended_space = run.text.endswith(" ")

    # Body-level paragraphs (doc.paragraphs = direct children of <w:body>)
    for para in doc.paragraphs:
        idx        = body_child_map.get(para._element, -1)
        is_sue_para = idx in single_ue_indices
        _apply_brackets(para, is_sue_para)

    # Table-cell paragraphs (never single-UE annotated at body level)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _apply_brackets(para, False)

    # ── Step 5: delete conditionally-excluded paragraphs ───────────────────
    # Each deleted paragraph also takes ONE adjacent empty spacer paragraph
    # with it (preferring the following one) so the blank-line separation
    # between the surviving neighbours stays single, not doubled.
    def _is_empty_para_el(el):
        if el.tag != _p_tag:
            return False
        return not "".join(
            t.text or "" for t in el.findall(f".//{qn('w:t')}")
        ).strip()

    _del_expanded = set(del_indices)
    for idx in sorted(del_indices):
        for adj in (idx + 1, idx - 1):
            if adj in _del_expanded:
                continue
            if 0 <= adj < len(body_children) and _is_empty_para_el(body_children[adj]):
                _del_expanded.add(adj)
                break

    for idx in sorted(_del_expanded, reverse=True):
        if idx < len(body_children):
            doc.element.body.remove(body_children[idx])

    # ── Step 6: font standardisation ───────────────────────────────────────
    cjk_font = FONT_CHINESE if language == "中文" else "Times New Roman"

    # Paragraphs whose text matches these patterns keep their bold intact.
    # Checked against the paragraph's full text AFTER substitution.
    _BOLD_KEEP_PATTERNS = [
        "Independent Service Auditor",   # AR section title
        "Management Assertion",          # MA section title (after co-name sub)
        "Management Statements",         # MA title (ISAE 3000/3402 templates)
        "Report of Its Assertions",      # MA title (SOC3 template)
        "Ernst & Young",                 # AR signature block (firm name)
        "Hua Ming",                      # AR signature block (firm name variant)
        "管理层认定",                     # CN MA section title (after co-name sub)
        "管理层声明",                     # CN MA title (ISAE 3000/3402 templates)
        "独立服务审计师",                 # CN AR title (…报告 / …鉴证报告)
        "安永华明",                       # CN AR signature block (firm name)
    ]
    # Date paragraph in MA: formatted date string is the entire paragraph text
    _DATE_RE = re.compile(
        r"^(?:[A-Z][a-z]+ \d{1,2}, \d{4}|\d{4}年\d{1,2}月\d{1,2}日)$"
    )
    # Company name used to detect the MA signature line (a short standalone para)
    _company_name_bold = subs.get("[Service organization name]", "")

    # Section titles get an explicit outline level so MA and AR appear as
    # 1st-level headings in Word's navigation pane (formatting unchanged).
    _TITLE_PATTERNS = [
        "Independent Service Auditor",
        "Management Assertion",
        "Management Statements",
        "Report of Its Assertions",
        "管理层认定",
        "管理层声明",
        "独立服务审计师",
    ]

    def _para_is_title(para):
        full = "".join(r.text for r in para.runs).strip()
        return len(full) <= 120 and any(p in full for p in _TITLE_PATTERNS)

    def _para_keep_bold(para):
        full = "".join(r.text for r in para.runs).strip()
        # Apply pattern matching only to short paragraphs (section titles,
        # signature-block lines). Without this guard the 10 000-char "We have
        # examined…" AR paragraph — which contains "Management Assertion" in
        # running body text — would be incorrectly made bold in its entirety.
        if len(full) <= 120 and any(pat in full for pat in _BOLD_KEEP_PATTERNS):
            return True
        if _DATE_RE.match(full):
            return True
        # MA signature line: standalone paragraph whose text is exactly the
        # service-organization name (e.g. "ABC Fintech Co., Ltd.")
        if _company_name_bold and full == _company_name_bold:
            return True
        # Short city / country line in the AR signature block
        # e.g. "Shanghai, China" or "中国 上海" (≤50 chars to avoid body text)
        if len(full) <= 50 and ("China" in full or "中国" in full):
            return True
        return False

    def _std_run(run, keep_bold=False):
        rPr = run._r.get_or_add_rPr()
        if not keep_bold:
            run.bold = False            # strip bold; italic is left untouched
            # Also explicitly disable complex-script bold (bCs) so that CJK
            # characters do not inherit bold from a heading paragraph style
            # after the document sections are merged.
            bCs = rPr.find(qn("w:bCs"))
            if bCs is None:
                bCs = OxmlElement("w:bCs")
                rPr.append(bCs)
            bCs.set(qn("w:val"), "0")
        else:
            run.bold = True             # ensure bold is explicitly set
        run.font.size = Pt(11)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"),    "Times New Roman")
        rFonts.set(qn("w:hAnsi"),    "Times New Roman")
        rFonts.set(qn("w:eastAsia"), cjk_font)
        rFonts.set(qn("w:cs"),       "Times New Roman")
        # Explicitly disable underline (overrides any inherited style underline)
        u_el = rPr.find(qn("w:u"))
        if u_el is None:
            u_el = OxmlElement("w:u")
            rPr.append(u_el)
        u_el.set(qn("w:val"), "none")

    def _clear_para_mark_bold(para):
        """Also strip bold from the paragraph-mark rPr (pPr/rPr)."""
        _pPr = para._p.find(qn("w:pPr"))
        if _pPr is None:
            return
        _pRPr = _pPr.find(qn("w:rPr"))
        if _pRPr is None:
            return
        for _bname in ("w:b", "w:bCs"):
            _bel = _pRPr.find(qn(_bname))
            if _bel is None:
                _bel = OxmlElement(_bname)
                _pRPr.append(_bel)
            _bel.set(qn("w:val"), "0")

    for para in doc.paragraphs:
        kb = _para_keep_bold(para)
        if _para_is_title(para):
            _set_outline_level(para, 0)
        for run in para.runs:
            _std_run(run, keep_bold=kb)
        if not kb:
            _clear_para_mark_bold(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    kb = _para_keep_bold(para)
                    for run in para.runs:
                        _std_run(run, keep_bold=kb)
                    if not kb:
                        _clear_para_mark_bold(para)

    # ── Step 7: strip spaces before punctuation ────────────────────────
    _PUNCT = '.,;:!?，。；：！？'

    def _strip_spaces_before_punct(para):
        for run in para.runs:
            if run.text:
                run.text = re.sub(r' +([' + re.escape(_PUNCT) + r'])', r'\1', run.text)
        active = [r for r in para.runs if r.text]
        for i in range(len(active) - 1):
            if active[i].text.endswith(' ') and active[i + 1].text[0] in _PUNCT:
                active[i].text = active[i].text.rstrip(' ')

    for para in doc.paragraphs:
        _strip_spaces_before_punct(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _strip_spaces_before_punct(para)

    buf = io.BytesIO()
    doc.save(buf)
    saved_bytes = buf.getvalue()

    # python-docx may silently drop <w:lvlOverride>/<w:startOverride> elements
    # when serialising the NumberingPart, which breaks lowerLetter list counters
    # (e.g. "a." / "b." become "l" / "l").  Re-inject the original numbering.xml
    # from the template to guarantee the counter-restart overrides are preserved.
    try:
        with zipfile.ZipFile(io.BytesIO(raw_bytes), "r") as _z_orig:
            if "word/numbering.xml" in _z_orig.namelist():
                _orig_num = _z_orig.read("word/numbering.xml")
                _buf_in  = io.BytesIO(saved_bytes)
                _buf_out = io.BytesIO()
                with zipfile.ZipFile(_buf_in, "r") as _z_in:
                    with zipfile.ZipFile(_buf_out, "w", zipfile.ZIP_DEFLATED) as _z_out:
                        for _info in _z_in.infolist():
                            _data = _z_in.read(_info.filename)
                            if _info.filename == "word/numbering.xml":
                                _data = _orig_num
                            _z_out.writestr(_info, _data)
                _buf_out.seek(0)
                return _buf_out.read()
    except Exception:
        pass
    return saved_bytes


def _remap_extra_numbering(base_bytes, extra_bytes):
    """
    Remap abstractNumId and numId values in extra_bytes so they do not clash
    with numbering definitions already present in base_bytes.

    Both documents' numbering.xml are inspected to find the highest IDs in
    the base; every ID in the extra is offset by that amount so no two
    definitions share an ID after the merge.

    Returns modified extra_bytes, or the original if either document has no
    numbering.xml or the extra has no numbered lists.
    """
    try:
        with zipfile.ZipFile(io.BytesIO(base_bytes), "r") as z:
            if "word/numbering.xml" not in z.namelist():
                return extra_bytes
            base_num_xml = z.read("word/numbering.xml").decode("utf-8")
    except Exception:
        return extra_bytes

    try:
        with zipfile.ZipFile(io.BytesIO(extra_bytes), "r") as z:
            if "word/numbering.xml" not in z.namelist():
                return extra_bytes
            extra_num_xml = z.read("word/numbering.xml").decode("utf-8")
            extra_doc_xml = z.read("word/document.xml").decode("utf-8")
    except Exception:
        return extra_bytes

    # Determine maximum IDs in base
    base_abs  = [int(x) for x in re.findall(r'w:abstractNumId="(\d+)"', base_num_xml)]
    base_nums = [int(x) for x in re.findall(r'<w:num\b[^>]*w:numId="(\d+)"', base_num_xml)]
    base_max_abs = max(base_abs)  if base_abs  else -1
    base_max_num = max(base_nums) if base_nums else  0

    # Collect IDs present in extra (process largest first to avoid partial hits)
    extra_abs  = sorted(
        set(int(x) for x in re.findall(r'<w:abstractNum\b[^>]*w:abstractNumId="(\d+)"', extra_num_xml)),
        reverse=True,
    )
    extra_nums = sorted(
        set(int(x) for x in re.findall(r'<w:num\b[^>]*w:numId="(\d+)"', extra_num_xml)),
        reverse=True,
    )

    if not extra_nums:
        return extra_bytes  # nothing numbered in extra

    abs_offset = base_max_abs + 1
    num_offset = base_max_num

    new_num_xml = extra_num_xml
    new_doc_xml = extra_doc_xml

    # Remap abstractNumId in numbering.xml (definition + cross-references)
    for old in extra_abs:
        new = old + abs_offset
        # <w:abstractNum w:abstractNumId="N"> — definition attribute
        new_num_xml = new_num_xml.replace(
            f'w:abstractNumId="{old}"', f'w:abstractNumId="{new}"'
        )
        # <w:abstractNumId w:val="N"/> — reference inside <w:num>
        new_num_xml = new_num_xml.replace(
            f'<w:abstractNumId w:val="{old}"', f'<w:abstractNumId w:val="{new}"'
        )

    # Remap numId in numbering.xml (definition) and in document.xml (references)
    for old in extra_nums:
        new = old + num_offset
        # <w:num w:numId="N"> — definition attribute
        new_num_xml = new_num_xml.replace(f'w:numId="{old}"', f'w:numId="{new}"')
        # <w:numId w:val="N"/> — paragraph numPr reference
        new_doc_xml = re.sub(
            rf'(<w:numId\s+w:val="){old}(")',
            rf'\g<1>{new}\g<2>',
            new_doc_xml,
        )

    # Write modified extra docx
    buf_in  = io.BytesIO(extra_bytes)
    buf_out = io.BytesIO()
    with zipfile.ZipFile(buf_in, "r") as zin:
        with zipfile.ZipFile(buf_out, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "word/numbering.xml":
                    data = new_num_xml.encode("utf-8")
                elif info.filename == "word/document.xml":
                    data = new_doc_xml.encode("utf-8")
                zout.writestr(info, data)
    buf_out.seek(0)
    return buf_out.read()


def _inject_numbering(merged_bytes, extra_bytes):
    """
    Append all abstractNum and num definitions from extra_bytes into the
    word/numbering.xml of merged_bytes.

    This is called after the document bodies have been merged so that the
    remapped numId values in the extra's paragraphs resolve to actual
    definitions in the merged file.
    """
    try:
        with zipfile.ZipFile(io.BytesIO(extra_bytes), "r") as z:
            if "word/numbering.xml" not in z.namelist():
                return merged_bytes
            extra_num_xml = z.read("word/numbering.xml").decode("utf-8")
    except Exception:
        return merged_bytes

    abs_blocks = re.findall(r"<w:abstractNum\b.*?</w:abstractNum>", extra_num_xml, re.DOTALL)
    num_blocks  = re.findall(r"<w:num\b.*?</w:num>",                extra_num_xml, re.DOTALL)

    if not abs_blocks and not num_blocks:
        return merged_bytes

    buf_in  = io.BytesIO(merged_bytes)
    buf_out = io.BytesIO()
    with zipfile.ZipFile(buf_in, "r") as zin:
        with zipfile.ZipFile(buf_out, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "word/numbering.xml":
                    num_xml = data.decode("utf-8")
                    # CT_Numbering requires all <w:abstractNum> elements to
                    # precede all <w:num> elements; appending both at the end
                    # interleaves them and Word flags the part as corrupt
                    # (its "repair" then scrambles the list definitions).
                    abs_inject = "\n".join(abs_blocks)
                    if abs_inject:
                        m = re.search(r"<w:num\b", num_xml)
                        pos = m.start() if m else num_xml.rfind("</w:numbering>")
                        num_xml = num_xml[:pos] + abs_inject + "\n" + num_xml[pos:]
                    num_inject = "\n".join(num_blocks)
                    if num_inject:
                        m = re.search(r"<w:numIdMacAtCleanup\b", num_xml)
                        pos = m.start() if m else num_xml.rfind("</w:numbering>")
                        num_xml = num_xml[:pos] + num_inject + "\n" + num_xml[pos:]
                    data = num_xml.encode("utf-8")
                zout.writestr(info, data)
    buf_out.seek(0)
    return buf_out.read()


def merge_docx_sections(*docs_bytes, split_sections=False):
    """
    Merge multiple docx byte strings into one document with page breaks between
    sections.

    Numbering IDs (abstractNumId / numId) in each extra document are remapped
    to avoid conflicts with the base document's numbering definitions, then the
    extra definitions are injected into the merged file's numbering.xml.  This
    preserves the original list styles (bullets stay bullets, alpha lists stay
    alpha lists) across section boundaries.

    When *split_sections* is True, each document boundary becomes a real Word
    *section break* (a paragraph carrying a clone of the base section
    properties) instead of a plain page break, so the merged document ends up
    with one Word section per input document.  This is what lets a single
    section (e.g. the Auditor's Report) carry its own header — see
    inject_ar_letterhead(), which overwrites that section's sectPr afterwards.
    The visual result is the same (each section still starts on a new page).
    """
    base_bytes = docs_bytes[0]

    # Remap numbering in each extra so its IDs don't clash with the base.
    # Use a rolling base: after remapping each extra, inject its definitions
    # into the running base so the next extra gets offsets that account for
    # all previously added abstractNum/num blocks, preventing duplicate IDs.
    extras_remapped = []
    running_base = base_bytes
    for eb in docs_bytes[1:]:
        remapped = _remap_extra_numbering(running_base, eb)
        extras_remapped.append(remapped)
        running_base = _inject_numbering(running_base, remapped)

    # Merge document bodies with python-docx
    base = Document(io.BytesIO(base_bytes))
    body = base.element.body
    last_sectPr = body.find(qn("w:sectPr"))
    # Source for cloned section-break properties (geometry/headers of the base
    # document). Each break that precedes an extra closes the *previous*
    # document's section; inject_ar_letterhead() later replaces the one section
    # that needs the EY letterhead.
    sect_props_src = last_sectPr

    for extra_bytes in extras_remapped:
        if split_sections and sect_props_src is not None:
            # Section break: empty paragraph whose pPr carries a clone of the
            # base section properties. Starts a new page like a page break, but
            # also closes the preceding document's Word section.
            p_br = OxmlElement("w:p")
            pPr  = OxmlElement("w:pPr")
            pPr.append(deepcopy(sect_props_src))
            p_br.append(pPr)
        else:
            # Plain page break before the next section
            p_br = OxmlElement("w:p")
            r_br = OxmlElement("w:r")
            br   = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            r_br.append(br)
            p_br.append(r_br)
        if last_sectPr is not None:
            body.insert(list(body).index(last_sectPr), p_br)
        else:
            body.append(p_br)

        extra = Document(io.BytesIO(extra_bytes))
        for elem in extra.element.body:
            if elem.tag != qn("w:sectPr"):
                if last_sectPr is not None:
                    body.insert(list(body).index(last_sectPr), deepcopy(elem))
                else:
                    body.append(deepcopy(elem))

    buf = io.BytesIO()
    base.save(buf)
    merged_bytes = buf.getvalue()

    # Inject remapped numbering definitions from each extra into the merged doc
    for extra_bytes in extras_remapped:
        merged_bytes = _inject_numbering(merged_bytes, extra_bytes)

    return merged_bytes


# Relationship / content-type strings for injected header parts
_REL_HEADER = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/header"
_CT_HEADER  = "application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"
_IMG_CONTENT_TYPES = {
    "png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
    "gif": "image/gif", "bmp": "image/bmp", "tiff": "image/tiff",
    "emf": "image/x-emf", "wmf": "image/x-wmf", "svg": "image/svg+xml",
}


def _pgmar_attr(sectpr, attr):
    """Return the integer value (in twips) of a <w:pgMar> attribute such as
    'top' or 'bottom' from a sectPr string, or None if absent."""
    m = re.search(r"<w:pgMar\b[^>]*>", sectpr)
    if not m:
        return None
    v = re.search(r'\bw:%s="(-?\d+)"' % attr, m.group(0))
    return int(v.group(1)) if v else None


def _set_pgmar_attr(sectpr, attr, value):
    """Return *sectpr* with the <w:pgMar> *attr* set to *value* (twips). If the
    attribute isn't present the sectPr is returned unchanged."""
    def repl(m):
        tag = m.group(0)
        return re.sub(r'\bw:%s="-?\d+"' % attr, 'w:%s="%d"' % (attr, value), tag)
    return re.sub(r"<w:pgMar\b[^>]*>", repl, sectpr, count=1)


def _extract_letterhead_parts(letterhead_path):
    """Read the header parts (and their images) + sectPr from a letterhead .docx.

    Returns (sectpr_xml, headers, support) where *headers* maps the header type
    ('default' / 'first') to {'xml', 'rels', 'images': {orig_target: bytes}} and
    *support* carries the letterhead's font and paragraph-style definitions
    ({'fonts': [<w:font> xml...], 'styles': [<w:style> xml...]}).  The fonts and
    styles must travel with the header: 'EYInterstate Light' is a named weight of
    the EYInterstate family, so without the matching font-table entry Word can't
    resolve the weighted name to the installed font and substitutes a default —
    even on a machine that has the font.  Footers are ignored (header-only)."""
    with zipfile.ZipFile(letterhead_path, "r") as z:
        names   = set(z.namelist())
        doc_xml = z.read("word/document.xml").decode("utf-8")
        m = re.search(r"<w:sectPr\b.*?</w:sectPr>", doc_xml, re.S)
        if not m:
            return None, {}, {"fonts": [], "styles": []}
        sectpr = m.group(0)

        support = {"fonts": [], "styles": []}
        if "word/fontTable.xml" in names:
            ft = z.read("word/fontTable.xml").decode("utf-8")
            support["fonts"] = re.findall(r"<w:font\b[^>]*?/>|<w:font\b.*?</w:font>", ft, re.S)
        if "word/styles.xml" in names:
            sx = z.read("word/styles.xml").decode("utf-8")
            support["styles"] = re.findall(r"<w:style\b.*?</w:style>", sx, re.S)

        rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
        rid_target = dict(re.findall(r'Id="([^"]+)"[^>]*?Target="([^"]+)"', rels))

        headers = {}
        for htype, rid in re.findall(
            r'<w:headerReference\s+w:type="([^"]+)"\s+r:id="([^"]+)"', sectpr
        ):
            target = rid_target.get(rid)
            if not target:
                continue
            part = "word/" + target.lstrip("./")
            if part not in names:
                continue
            hxml    = z.read(part)
            relname = "word/_rels/" + target.lstrip("./") + ".rels"
            hrels   = z.read(relname).decode("utf-8") if relname in names else (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<Relationships xmlns="http://schemas.openxmlformats.org/'
                'package/2006/relationships"></Relationships>'
            )
            images = {}
            for _iid, itarget in re.findall(r'Id="([^"]+)"[^>]*?Target="([^"]+)"', hrels):
                ipart = "word/" + itarget.lstrip("./")
                if ipart in names:
                    images[itarget] = z.read(ipart)
            headers[htype] = {"xml": hxml, "rels": hrels, "images": images}
        return sectpr, headers, support


def inject_ar_letterhead(docx_bytes, letterhead_path, ar_index):
    """Give the AR section of a merged document its own EY letterhead header.

    Copies the letterhead's header part(s) + logo image into the merged package
    and overwrites the ar_index-th sectPr (0-based, in document order) with the
    letterhead's section geometry + header references, so the Auditor's Report
    pages — and only those — carry the letterhead (full letterhead on the first
    page via the 'first' header, continuation header on later pages). Footers
    are not injected. Requires the document to have been merged with
    split_sections=True so one sectPr corresponds to the AR section."""
    sectpr_src, headers, support = _extract_letterhead_parts(letterhead_path)
    if not sectpr_src or not headers:
        return docx_bytes

    buf_in  = io.BytesIO(docx_bytes)
    buf_out = io.BytesIO()
    with zipfile.ZipFile(buf_in, "r") as zin:
        names    = set(zin.namelist())
        doc_xml  = zin.read("word/document.xml").decode("utf-8")
        rels_xml = zin.read("word/_rels/document.xml.rels").decode("utf-8")
        ct_xml   = zin.read("[Content_Types].xml").decode("utf-8")

        # Transplant the letterhead's font declarations and paragraph styles
        # that the merged document is missing, so the header runs resolve
        # 'EYInterstate Light' (a named weight) to the installed font instead
        # of being substituted. Only names/ids not already present are added.
        def _missing(blocks, id_attr, have):
            out = []
            for blk in blocks:
                m = re.search(id_attr + r'="([^"]+)"', blk)
                if m and m.group(1) not in have:
                    have.add(m.group(1))   # also de-dupes within the added set
                    out.append(blk)
            return out

        font_table_xml = None
        if "word/fontTable.xml" in names and support["fonts"]:
            ft   = zin.read("word/fontTable.xml").decode("utf-8")
            have = set(re.findall(r'<w:font\b[^>]*\bw:name="([^"]+)"', ft))
            add  = _missing(support["fonts"], "w:name", have)
            if add and "</w:fonts>" in ft:
                font_table_xml = ft.replace("</w:fonts>", "".join(add) + "</w:fonts>")

        # Isolate the letterhead header's style chain from the merged document.
        # The merged doc is based on the MA template, whose 'Normal' style is
        # bold and whose 'Header' style has a bottom border. The letterhead's
        # header paragraphs reference styles by id (Header, EYBusinessaddress, …
        # → basedOn 'Normal'); sharing those ids with the host makes the header
        # inherit the wrong weight and a stray black line. So copy every style
        # the header depends on (following basedOn/link/next) under a private id
        # namespace, repoint the copies' internal references, and rewrite the
        # header XML (in the loop below) to use the private ids — the letterhead
        # header then renders exactly as authored, independent of the host's
        # styles. style_id_map (orig id -> private id) drives the header rewrite.
        style_id_map = {}
        styles_xml = None
        if "word/styles.xml" in names and support["styles"]:
            sx       = zin.read("word/styles.xml").decode("utf-8")
            existing = set(re.findall(r'w:styleId="([^"]+)"', sx))

            lh_style = {}
            for blk in support["styles"]:
                m = re.search(r'w:styleId="([^"]+)"', blk)
                if m:
                    lh_style[m.group(1)] = blk

            seed = set()
            for h in headers.values():
                hx = h["xml"].decode("utf-8") if isinstance(h["xml"], bytes) else h["xml"]
                seed.update(re.findall(r'<w:(?:pStyle|rStyle|tblStyle) w:val="([^"]+)"', hx))

            # Transitive closure over basedOn/link/next, restricted to styles the
            # letterhead actually defines (a ref outside lh_style — e.g. a bare
            # built-in id — is left to resolve against the host as a last resort).
            closure, stack = set(), list(seed)
            while stack:
                sid = stack.pop()
                if sid in closure or sid not in lh_style:
                    continue
                closure.add(sid)
                stack.extend(re.findall(
                    r'<w:(?:basedOn|link|next) w:val="([^"]+)"', lh_style[sid]))

            def _uniq(name):
                cand, i = name, 1
                while cand in existing:
                    cand, i = "%s%d" % (name, i), i + 1
                existing.add(cand)
                return cand
            style_id_map = {sid: _uniq("LH" + sid) for sid in closure}

            def _remap_ref(m):
                return m.group(1) + style_id_map.get(m.group(2), m.group(2)) + m.group(3)

            new_blocks = []
            for sid, new_id in style_id_map.items():
                blk = re.sub(r'(w:styleId=")[^"]+(")',
                             lambda m, _i=new_id: m.group(1) + _i + m.group(2),
                             lh_style[sid], count=1)
                blk = re.sub(r'(<w:(?:basedOn|link|next) w:val=")([^"]+)(")',
                             _remap_ref, blk)
                new_blocks.append(blk)

            if new_blocks and "</w:styles>" in sx:
                styles_xml = sx.replace("</w:styles>", "".join(new_blocks) + "</w:styles>")

        # Allocate rIds that don't clash with the existing document rels
        used    = [int(n) for n in re.findall(r'Id="rId(\d+)"', rels_xml)]
        next_id = (max(used) + 1) if used else 1

        img_map      = {}   # orig target -> new media target (relative to word/)
        media_writes = {}   # new word/media/... path -> bytes
        header_parts = {}   # htype -> {part, relpart, rid, xml, rels(bytes)}
        new_rels_entries = []
        new_ct_overrides = []
        img_n = 1

        for htype, h in headers.items():
            new_rels = h["rels"]
            for orig_target, data in h["images"].items():
                if orig_target not in img_map:
                    ext = (os.path.splitext(orig_target)[1] or ".png")
                    new_media = "media/lh_letterhead_img%d%s" % (img_n, ext)
                    img_map[orig_target] = new_media
                    media_writes["word/" + new_media] = data
                    img_n += 1
                for variant in (orig_target, "./" + orig_target):
                    new_rels = new_rels.replace(
                        'Target="%s"' % variant, 'Target="%s"' % img_map[orig_target]
                    )
            # Repoint the header's style references to the private copies added
            # above, so it uses the letterhead's own (non-bold, borderless)
            # styles rather than the host document's same-named ones.
            hxml = h["xml"]
            if style_id_map:
                hs = hxml.decode("utf-8") if isinstance(hxml, bytes) else hxml
                hs = re.sub(
                    r'(<w:(?:pStyle|rStyle|tblStyle) w:val=")([^"]+)(")',
                    lambda m: m.group(1) + style_id_map.get(m.group(2), m.group(2)) + m.group(3),
                    hs,
                )
                hxml = hs.encode("utf-8")
            base_name = "lh_header_%s.xml" % htype
            rid       = "rId%d" % next_id
            next_id  += 1
            header_parts[htype] = {
                "part":    "word/" + base_name,
                "relpart": "word/_rels/" + base_name + ".rels",
                "rid":     rid,
                "xml":     hxml,
                "rels":    new_rels.encode("utf-8"),
            }
            new_rels_entries.append(
                '<Relationship Id="%s" Type="%s" Target="%s"/>'
                % (rid, _REL_HEADER, base_name)
            )
            new_ct_overrides.append(
                '<Override PartName="/word/%s" ContentType="%s"/>' % (base_name, _CT_HEADER)
            )

        # Build the AR section properties from the letterhead's sectPr:
        # drop footers, repoint the header references to the injected parts.
        new_sectpr = re.sub(r"<w:footerReference\b[^>]*/>", "", sectpr_src)

        def _repoint(m):
            htype = m.group(1)
            if htype in header_parts:
                return '<w:headerReference w:type="%s" r:id="%s"/>' % (
                    htype, header_parts[htype]["rid"])
            return ""   # a header type we couldn't extract — drop the reference

        new_sectpr = re.sub(
            r'<w:headerReference\s+w:type="([^"]+)"\s+r:id="[^"]+"\s*/>',
            _repoint, new_sectpr,
        )

        sectprs = list(re.finditer(r"<w:sectPr\b.*?</w:sectPr>", doc_xml, re.S))
        if ar_index >= len(sectprs):
            return docx_bytes   # structure unexpected — leave document untouched

        # The letterhead's sectPr replaces the whole AR page geometry. We only
        # want its enlarged top margin (to clear the banner); the side margins
        # should stay as the AR template authored them (~1800), so the AR body
        # text keeps the same width/indent as the rest of the report rather than
        # inheriting the letterhead's own narrow side margins (e.g. Shanghai
        # 1368/1282). Restore the AR section's original left/right (and gutter)
        # margins onto the letterhead-derived sectPr. This runs for both output
        # modes — complete report and "MA + AR only" — since both build the AR
        # page through this same function.
        orig_ar = sectprs[ar_index].group(0)
        for _side in ("left", "right", "gutter"):
            _orig = _pgmar_attr(orig_ar, _side)
            if _orig is not None:
                new_sectpr = _set_pgmar_attr(new_sectpr, _side, _orig)

        # The letterhead also carries its own (small) bottom margin, which leaves
        # the last line uncomfortably close to the page edge. Keep the AR
        # section's bottom margin close to the EY template's original value (the
        # one that looks fine without a letterhead) — allow at most a 10%
        # reduction — so there's reasonable breathing room at the page bottom.
        orig_bottom = _pgmar_attr(orig_ar, "bottom")
        lh_bottom   = _pgmar_attr(new_sectpr, "bottom")
        if orig_bottom and lh_bottom is not None:
            min_bottom = int(round(orig_bottom * 0.9))
            if lh_bottom < min_bottom:
                new_sectpr = _set_pgmar_attr(new_sectpr, "bottom", min_bottom)

        # Word propagates a section's header forward: a later section with no
        # headerReference of its own inherits the previous section's header. So
        # the AR letterhead would bleed onto the report body after it. Give every
        # section *after* AR an explicit blank header (for the types it doesn't
        # already define) to break that inheritance.
        post = list(range(ar_index + 1, len(sectprs)))
        blank_part = None
        if post:
            blank_rid  = "rId%d" % next_id
            next_id   += 1
            blank_part = "lh_header_blank.xml"
            new_rels_entries.append(
                '<Relationship Id="%s" Type="%s" Target="%s"/>'
                % (blank_rid, _REL_HEADER, blank_part)
            )
            new_ct_overrides.append(
                '<Override PartName="/word/%s" ContentType="%s"/>' % (blank_part, _CT_HEADER)
            )

        replacements = {ar_index: new_sectpr}
        for i in post:
            seg = sectprs[i].group(0)
            mo  = re.match(r"<w:sectPr\b[^>]*>", seg)
            if not mo:
                continue
            have_types = set(re.findall(r'<w:headerReference\s+w:type="([^"]+)"', seg))
            adds = "".join(
                '<w:headerReference w:type="%s" r:id="%s"/>' % (htype, blank_rid)
                for htype in ("default", "first") if htype not in have_types
            )
            if adds:
                replacements[i] = mo.group(0) + adds + seg[mo.end():]

        # Splice replacements in right-to-left so earlier match offsets stay valid
        for i in sorted(replacements, reverse=True):
            sp = sectprs[i]
            doc_xml = doc_xml[:sp.start()] + replacements[i] + doc_xml[sp.end():]

        # Register the new relationships, content types, and image defaults
        rels_xml = rels_xml.replace(
            "</Relationships>", "".join(new_rels_entries) + "</Relationships>"
        )
        for media_path in media_writes:
            ext = os.path.splitext(media_path)[1].lstrip(".").lower()
            if ext and ('Extension="%s"' % ext) not in ct_xml:
                ct = _IMG_CONTENT_TYPES.get(ext, "application/octet-stream")
                ct_xml = ct_xml.replace(
                    "</Types>", '<Default Extension="%s" ContentType="%s"/></Types>' % (ext, ct)
                )
        ct_xml = ct_xml.replace("</Types>", "".join(new_ct_overrides) + "</Types>")

        with zipfile.ZipFile(buf_out, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                fn = info.filename
                if fn == "word/document.xml":
                    data = doc_xml.encode("utf-8")
                elif fn == "word/_rels/document.xml.rels":
                    data = rels_xml.encode("utf-8")
                elif fn == "[Content_Types].xml":
                    data = ct_xml.encode("utf-8")
                elif fn == "word/fontTable.xml" and font_table_xml is not None:
                    data = font_table_xml.encode("utf-8")
                elif fn == "word/styles.xml" and styles_xml is not None:
                    data = styles_xml.encode("utf-8")
                else:
                    data = zin.read(fn)
                zout.writestr(info, data)
            for hp in header_parts.values():
                zout.writestr(hp["part"], hp["xml"])
                zout.writestr(hp["relpart"], hp["rels"])
            for path, data in media_writes.items():
                zout.writestr(path, data)
            if blank_part:
                zout.writestr(
                    "word/" + blank_part,
                    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                    '<w:hdr xmlns:w="http://schemas.openxmlformats.org/'
                    'wordprocessingml/2006/main"><w:p/></w:hdr>',
                )

    buf_out.seek(0)
    return buf_out.read()


def strip_page_top_empty_paragraphs(docx_bytes):
    """Remove blank paragraphs at the top of each page that begins
    *deterministically* — the document start and immediately after an explicit
    page break or section break (e.g. the blank lines an EY template left at the
    top of the Auditor's Report to make room for a letterhead, now redundant
    once the letterhead sits in the header + page margin).

    Soft page breaks (where text naturally flows onto the next page) are decided
    by Word at render time and are not recorded in the file, so blank lines at
    the top of those pages cannot be detected here — only break-initiated page
    tops can. Operates on body-level paragraphs only; the break/section
    paragraphs themselves are preserved (they carry the page break / sectPr)."""
    doc  = Document(io.BytesIO(docx_bytes))
    body = doc.element.body
    paras = body.findall(qn("w:p"))

    def _has_sectpr(p):
        pPr = p.find(qn("w:pPr"))
        return pPr is not None and pPr.find(qn("w:sectPr")) is not None

    def _has_pagebreak(p):
        return any(br.get(qn("w:type")) == "page" for br in p.iter(qn("w:br")))

    def _is_break(p):
        return _has_sectpr(p) or _has_pagebreak(p)

    def _is_blank(p):
        if _has_sectpr(p) or _has_pagebreak(p):
            return False
        for tag in ("w:drawing", "w:pict", "w:object"):
            if p.find(".//" + qn(tag)) is not None:
                return False
        return "".join(t.text or "" for t in p.iter(qn("w:t"))).strip() == ""

    to_remove = []
    # Leading blanks at the very start of the document
    i = 0
    while i < len(paras) and _is_blank(paras[i]):
        to_remove.append(paras[i])
        i += 1
    # Blanks immediately after each explicit page/section break
    for k, p in enumerate(paras):
        if _is_break(p):
            j = k + 1
            while j < len(paras) and _is_blank(paras[j]):
                to_remove.append(paras[j])
                j += 1

    for p in to_remove:
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def enforce_line_spacing(docx_bytes, spacing=1.15):
    """
    Final pass over the finished document: set every paragraph (body and
    table cells, including nested tables) to the given multiple line spacing,
    and normalise every run to Times New Roman Latin/complex-script fonts and
    black text, and the CJK (eastAsia) font to 黑体 — so the whole document
    is uniform (Times New Roman for Latin, 黑体 for Chinese) regardless of
    what each merged section set.

    Re-injects the original numbering.xml afterwards because python-docx may
    silently drop <w:lvlOverride>/<w:startOverride> elements on save (same
    workaround as in fill_and_process_template).
    """
    doc = Document(io.BytesIO(docx_bytes))

    def _walk(paragraphs, tables):
        for para in paragraphs:
            pf = para.paragraph_format
            pf.line_spacing = spacing
            pf.space_before = Pt(0)
            pf.space_after  = Pt(0)
            # beforeLines/afterLines and autospacing attributes take precedence
            # over before/after in Word — strip them so 0 actually applies.
            sp_el = para._p.pPr.find(qn("w:spacing"))
            if sp_el is not None:
                for attr in ("w:beforeLines", "w:afterLines",
                             "w:beforeAutospacing", "w:afterAutospacing"):
                    if sp_el.get(qn(attr)) is not None:
                        del sp_el.attrib[qn(attr)]
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    _walk(cell.paragraphs, cell.tables)

    _walk(doc.paragraphs, doc.tables)

    # Iterate raw <w:r> elements so runs inside hyperlinks, text boxes and
    # nested tables — which para.runs does not expose — are covered too.
    # get_or_add_* keeps rPr children in schema order (avoids Word repair).
    for r_el in doc.element.body.iter(qn("w:r")):
        rPr = r_el.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:ascii"),    "Times New Roman")
        rFonts.set(qn("w:hAnsi"),    "Times New Roman")
        rFonts.set(qn("w:cs"),       "Times New Roman")
        rFonts.set(qn("w:eastAsia"), FONT_CHINESE)
        for theme_attr in ("w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:eastAsiaTheme"):
            if rFonts.get(qn(theme_attr)) is not None:
                del rFonts.attrib[qn(theme_attr)]
        color = rPr.get_or_add_color()
        color.set(qn("w:val"), "000000")
        for theme_attr in ("w:themeColor", "w:themeTint", "w:themeShade"):
            if color.get(qn(theme_attr)) is not None:
                del color.attrib[qn(theme_attr)]
        # Uniform 11 pt body text everywhere (Latin + complex script)
        rPr.sz_val = Pt(11)
        szCs = rPr.find(qn("w:szCs"))
        if szCs is None:
            szCs = OxmlElement("w:szCs")
            rPr.find(qn("w:sz")).addnext(szCs)
        szCs.set(qn("w:val"), "22")
        # EY template styles define Normal as BOLD; the Dify-generated sections
        # merged into the template document inherit it and render whole-bold.
        # Pin bold/underline OFF wherever the run does not set them itself, so
        # explicit bold (titles, **bold** markdown, table headers) and explicit
        # underline (H3 headings) are preserved.
        b_el = rPr.find(qn("w:b"))
        if b_el is None:
            b_el = OxmlElement("w:b")
            b_el.set(qn("w:val"), "0")
            rFonts.addnext(b_el)
        if rPr.find(qn("w:bCs")) is None:
            bcs_el = OxmlElement("w:bCs")
            bcs_el.set(qn("w:val"), "0")
            b_el.addnext(bcs_el)
        if rPr.find(qn("w:u")) is None:
            u_el = OxmlElement("w:u")
            u_el.set(qn("w:val"), "none")
            rPr.append(u_el)

    # Paragraph-mark rPr (w:pPr/w:rPr) controls how list numbers/bullets are
    # rendered when the numbering level itself names no font — e.g. the CN AR
    # "a." items showed in 黑体 because the mark carried eastAsia=黑体 and no
    # ascii font. Force Times New Roman Latin + 华文楷体 CJK + black here too.
    _mark_tail_tags = (qn("w:sectPr"), qn("w:pPrChange"))
    for pPr in doc.element.body.iter(qn("w:pPr")):
        rPr = pPr.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            tail = next((ch for ch in pPr if ch.tag in _mark_tail_tags), None)
            if tail is not None:
                tail.addprevious(rPr)
            else:
                pPr.append(rPr)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rStyle = rPr.find(qn("w:rStyle"))
            if rStyle is not None:
                rStyle.addnext(rFonts)
            else:
                rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"),    "Times New Roman")
        rFonts.set(qn("w:hAnsi"),    "Times New Roman")
        rFonts.set(qn("w:cs"),       "Times New Roman")
        rFonts.set(qn("w:eastAsia"), FONT_CHINESE)
        for theme_attr in ("w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:eastAsiaTheme"):
            if rFonts.get(qn(theme_attr)) is not None:
                del rFonts.attrib[qn(theme_attr)]
        color = rPr.find(qn("w:color"))
        if color is None:
            color = OxmlElement("w:color")
            rFonts.addnext(color)
        color.set(qn("w:val"), "000000")
        for theme_attr in ("w:themeColor", "w:themeTint", "w:themeShade"):
            if color.get(qn(theme_attr)) is not None:
                del color.attrib[qn(theme_attr)]
        # List number glyphs and empty-paragraph line heights at 11 pt as well
        rPr.sz_val = Pt(11)
        mszCs = rPr.find(qn("w:szCs"))
        if mszCs is None:
            mszCs = OxmlElement("w:szCs")
            rPr.find(qn("w:sz")).addnext(mszCs)
        mszCs.set(qn("w:val"), "22")
        # Keep list number glyphs non-bold unless the mark sets bold itself
        # (the bold Normal style would otherwise bleed into them).
        mb_el = rPr.find(qn("w:b"))
        if mb_el is None:
            mb_el = OxmlElement("w:b")
            mb_el.set(qn("w:val"), "0")
            rFonts.addnext(mb_el)
        if rPr.find(qn("w:bCs")) is None:
            mbcs_el = OxmlElement("w:bCs")
            mbcs_el.set(qn("w:val"), "0")
            mb_el.addnext(mbcs_el)

    buf = io.BytesIO()
    doc.save(buf)
    saved_bytes = buf.getvalue()

    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as z_orig:
            if "word/numbering.xml" not in z_orig.namelist():
                return saved_bytes
            orig_num = z_orig.read("word/numbering.xml")
        buf_in, buf_out = io.BytesIO(saved_bytes), io.BytesIO()
        with zipfile.ZipFile(buf_in, "r") as z_in:
            with zipfile.ZipFile(buf_out, "w", zipfile.ZIP_DEFLATED) as z_out:
                for info in z_in.infolist():
                    data = z_in.read(info.filename)
                    if info.filename == "word/numbering.xml":
                        data = orig_num
                    z_out.writestr(info, data)
        buf_out.seek(0)
        return buf_out.read()
    except Exception:
        return saved_bytes


def build_substitutions(ui, tc):
    """Build the EN + CN placeholder → value substitution dict."""
    language      = ui.get("Output_language", "English")
    company_name  = ui.get("Company_name", "")
    co_short_name = ui.get("Co_short_name", "")
    system_name   = ui.get("System_or_service_name", "")
    # EN reports: fix names typed in all-lowercase / ALL-CAPS (mixed case kept)
    if language != "中文":
        company_name  = _capitalize_name(company_name)
        co_short_name = _capitalize_name(co_short_name)
        system_name   = _capitalize_name(system_name)
    report_type   = ui.get("Report_type", "")
    subservice_org = ui.get("Subservice_org", "")
    # City typed in lowercase ('shanghai') → 'Shanghai'; CJK ('上海') untouched
    signing_city  = _capitalize_name(tc.get("signing_city", ""))

    # Format raw YYYY-MM-DD dates to "Month D, YYYY" / "YYYY年M月D日"
    period_start = _format_date(ui.get("Period_start", ""), language)
    period_end   = _format_date(ui.get("Period_end",   ""), language)
    report_date  = _format_date(tc.get("report_date",  ""), language)

    # Determine date placeholders based on report type
    if "TYPE1" in report_type:
        single_date = period_start  # Type I: "as of" date
    else:
        single_date = period_end    # Type II: period end date

    period_str = (
        f"{period_start} to {period_end}"
        if period_start and period_end
        else period_start or period_end
    )
    period_str_cn = (
        f"{period_start}至{period_end}"
        if period_start and period_end
        else period_start or period_end
    )

    # Parse SSO entries from subservice_org.
    # Supported formats (one org per line):
    #   "Name | Short Name | Services"  (preferred — 3 columns)
    #   "Name | Services"               (2 columns — short name defaults to full name)
    #   "Name"                          (1 column — both short name and services empty)
    _sso_entries = []  # list of (name, short_name, services)
    if subservice_org:
        for _raw in subservice_org.strip().splitlines():
            _ln = _raw.strip()
            if not _ln:
                continue
            _parts = [p.strip() for p in _ln.split("|")]
            if len(_parts) >= 3:
                _sso_entries.append((_parts[0], _parts[1], "|".join(_parts[2:]).strip()))
            elif len(_parts) == 2:
                _sso_entries.append((_parts[0], _parts[0], _parts[1]))
            else:
                _sso_entries.append((_parts[0], _parts[0], ""))
    _sso_a = _sso_entries[0] if len(_sso_entries) > 0 else ("", "", "")
    _sso_b = _sso_entries[1] if len(_sso_entries) > 1 else ("", "", "")
    _sso_c = _sso_entries[2] if len(_sso_entries) > 2 else ("", "", "")
    sso_name       = _sso_a[0]
    sso_short_name = _sso_a[1]
    sso_services   = _sso_a[2]

    # Addressee line: replace the combined "Management of/Board of Directors of"
    # placeholder before the generic [Service organization name] sub runs.
    addressee = tc.get("addressee_choice", "Management")
    if addressee == "Board of Directors":
        addr_label = "Board of Directors"
        addr_cn = "董事会"
    else:
        addr_label = "Management"
        addr_cn = "管理层"
    subs = {
        # Addressee line — must come BEFORE the generic [Service organization name] sub.
        # Templates use two EN wordings for the combined placeholder.
        "To the Management of/Board of Directors of [Service organization name]":
            f"To the {addr_label} of {company_name}",
        "To the Management/Board of Directors of [Service organization name]":
            f"To the {addr_label} of {company_name}",
        # CN combined addressee: 董事会/管理层 → the chosen one
        "董事会/管理层": addr_cn,
        # EN placeholders
        "[Service organization name]":          company_name,
        "[Service organization short name]":     co_short_name,
        "[Service organization\u2019s system]":  system_name,   # right single quote
        "[Service organization's system]":       system_name,   # straight apostrophe
        "[type or name of system]":              system_name,
        # Note: [or identification of the function performed by the System] is
        # handled by the [or ..] regex removal in fill_and_process_template.
        "[date] to [date]":                      period_str,
        "[date]":                                single_date,
        "[Date of the service auditor\u2019s report]": report_date,  # right quote
        "[Date of the service auditor's report]":      report_date,  # straight quote
        "[Date of report]":                      report_date,
        # AR signature: the bracketed alternative firm name is dropped — the
        # "<firm> <city> Branch" line above it is kept (city adjusted below).
        "[Ernst & Young Hua Ming LLP]":          "",
        # City: replace the whole "default_city[alternatives]" pattern
        "Shanghai[Beijing, Shenzhen]":           signing_city,
        "[Beijing, Shenzhen]":                   signing_city,
        # Single-SSO EN placeholders (generic, appear in non-A/B templates)
        "[Subservice organization name]":        sso_name,
        "[Subservice organization short name]":  sso_short_name,
        "[identify the function or service provided by the subservice organization]": sso_services,
        "[description of services provided]":    sso_services,
        # Multi-SSO EN placeholders (A / B / C variants in templates like WP 10.1)
        "[Subservice organization A name]":      _sso_a[0],
        "[Subservice organization A short name]": _sso_a[1],
        "[identify the function or service provided by the subservice organization A]": _sso_a[2],
        "[Subservice organization B name]":      _sso_b[0],
        "[Subservice organization B short name]": _sso_b[1],
        "[identify the function or service provided by the subservice organization B]": _sso_b[2],
        "[Subservice organization C short name]": _sso_c[1],
        # Capital-O variant found in some templates
        "[Service Organization short name]":     co_short_name,
        "[Company name]":                        company_name,
        "[Service System name]":                 system_name,
        "[V]":                                   "V",
        # CN placeholders — templates use fullwidth lenticular brackets
        # 【】 (U+3010/U+3011)
        "\u3010\u670d\u52a1\u673a\u6784\u540d\u79f0\u3011": company_name,   # 【服务机构名称】
        "\u3010\u670d\u52a1\u673a\u6784\u7b80\u79f0\u3011": co_short_name,  # 【服务机构简称】
        "\u3010\u670d\u52a1\u673a\u6784\u4f53\u7cfb\u540d\u79f0\u3011": system_name,  # 【服务机构体系名称】
        "【服务机构服务体系名称】": system_name,  # 【服务机构服务体系名称】 (SOC3 CN variant)
        # Single-SSO CN placeholders
        "【子服务机构名称】":              sso_name,
        "【子服务机构简称】":              sso_short_name,
        "【子服务机构】":                  sso_name,
        "【子服务机构提供的功能或服务】":  sso_services,
        "【子服务机构的服务类型或内容】":  sso_services,
        # Multi-SSO CN placeholders (A / B / C variants)
        "【子服务机构A名称】":             _sso_a[0],
        "【子服务机构A简称】":             _sso_a[1],
        "【子服务机构A的服务类型或内容】": _sso_a[2],
        "【子服务机构B名称】":             _sso_b[0],
        "【子服务机构B简称】":             _sso_b[1],
        "【子服务机构B的服务类型或内容】": _sso_b[2],
        "【子服务机构C简称】":             _sso_c[1],
        # CN period range — must come BEFORE the single 【日期】 sub below
        # (templates write Type II periods as 自【日期】至【日期】止).
        "【日期】至【日期】": period_str_cn,
        "\u3010\u65e5\u671f\u3011": single_date,   # 【日期】
        "\u3010\u62a5\u544a\u65e5\u3011": report_date,  # 【报告日】
        # City CN: replace the default+alternatives pattern
        "\u4e2d\u56fd \u4e0a\u6d77\u3010\u6216\u4e2d\u56fd \u5317\u4eac\u6216\u4e2d\u56fd \u6df1\u5733\u3011": f"\u4e2d\u56fd {signing_city}",  # 中国 上海【或中国 北京或中国 深圳】→ 中国 {city}
        "\u3010\u6216\u5b89\u6c38\u534e\u660e\u4f1a\u8ba1\u5e08\u4e8b\u52a1\u6240\uff08\u7279\u6b8a\u666e\u901a\u5408\u4f19\uff09\u3011": "",  # 【或安永华明...】→ empty (keep branch)
    }
    # AR signature branch line: adjust "… Shanghai Branch" / "…上海分所" to the
    # signing city. Skipped when the city IS Shanghai — substituting a value
    # identical to its placeholder would loop forever in the cross-run pass.
    if signing_city and signing_city not in ("Shanghai", "上海"):
        subs["Ernst & Young Hua Ming LLP Shanghai Branch"] = (
            f"{EY_FIRM_NAME} {signing_city} Branch"
        )
        subs["安永华明会计师事务所（特殊普通合伙）上海分所"] = (
            f"安永华明会计师事务所（特殊普通合伙）{signing_city}分所"
        )
    # Trust Service Criteria — fill the bracketed "Relevant to [Security, …]"
    # placeholder from the user's TSC selection (SOC2). Only the BRACKETED
    # variants are substituted; the unbracketed official criteria title
    # ("…Trust Services Criteria for Security, Availability, …") must stay.
    _TSC_DEFS = [
        ("is_Security",             "Security",             "安全性"),
        ("is_Availability",         "Availability",         "可用性"),
        ("is_Processing_Integrity", "Processing Integrity", "进程完整性"),
        ("is_Confidentiality",      "Confidentiality",      "保密性"),
        ("is_Privacy",              "Privacy",              "隐私性"),
    ]
    tsc_en = [en for key, en, _cn in _TSC_DEFS if ui.get(key)]
    tsc_cn = [cn for key, _en, cn in _TSC_DEFS if ui.get(key)]
    if tsc_en:
        if len(tsc_en) == 1:
            tsc_en_str = tsc_en[0]
        elif len(tsc_en) == 2:
            tsc_en_str = f"{tsc_en[0]} and {tsc_en[1]}"
        else:
            tsc_en_str = ", ".join(tsc_en[:-1]) + f", and {tsc_en[-1]}"
        if len(tsc_cn) == 1:
            tsc_cn_str = tsc_cn[0]
        else:
            tsc_cn_str = "、".join(tsc_cn[:-1]) + f"以及{tsc_cn[-1]}"
        tsc_en_lower = tsc_en_str.lower()
        # EN templates use both Oxford-comma and non-Oxford variants, plus a
        # second LOWERCASE occurrence ("…trust services criteria relevant to
        # [security, …]") with its own wording variants.
        subs["[Security, Availability, Processing Integrity, Confidentiality, and Privacy]"] = tsc_en_str
        subs["[Security, Availability, Processing Integrity, Confidentiality and Privacy]"]  = tsc_en_str
        subs["[security, availability, processing integrity, confidentiality, and privacy]"] = tsc_en_lower
        subs["[security, availability, processing integrity, confidentiality, privacy]"]     = tsc_en_lower
        subs["[security, availability, processing integrity and confidentiality, privacy]"]  = tsc_en_lower
        # CN templates: 以及 / 及 variants plus one with a 进程性、完整性 typo
        subs["【安全性、可用性、进程完整性、保密性以及隐私性】"]   = tsc_cn_str
        subs["【安全性、可用性、进程完整性、保密性及隐私性】"]     = tsc_cn_str
        subs["【安全性、可用性、进程性、完整性、保密性以及隐私性】"] = tsc_cn_str
    # When transaction-processing wording is excluded, remove the inline
    # transaction phrases and substitute "[or identification of the function
    # performed by the System]" with the user-supplied system function description.
    if not tc.get("has_transaction_processing", True):
        sys_fn = ui.get("Systems_function", "")
        # EN variants — remove transaction-processing phrases
        subs["for processing user entities\u2019 transactions"] = ""   # right-quote
        subs["for processing user entities' transactions"] = ""        # straight-quote
        subs["for processing their transactions"] = ""
        # "in processing or reporting transactions" → "in" so that the following
        # [or identification...] substitution produces "in <system function>"
        # rather than "in processing or reporting <system function>" (issue 8)
        subs["in processing or reporting transactions"] = "in"
        subs["[or identification of the function performed by the System]"] = sys_fn
        subs["[or identification of the function performed by the system]"] = sys_fn
        # CN counterpart — without this sub the generic 【或…】 removal would
        # delete the alternative instead of filling in the system function.
        subs["【或确定体系执行的功能】"] = sys_fn
        # Remove the "auditors" clause from the "intended solely for…" paragraph.
        # That clause only applies when the system processes user-entity transactions.
        # Both right-quote (U+2019) and straight-apostrophe variants are covered.
        subs[
            ", and their auditors who audit and report on such user entities\u2019 "
            "financial statements or internal control over financial reporting"
        ] = ""
        subs[
            ", and their auditors who audit and report on such user entities' "
            "financial statements or internal control over financial reporting"
        ] = ""
    return subs


def build_flags(tc):
    """Build the boolean deletion-flag dict from template_config."""
    return {
        "cuec_identified":            tc.get("cuec_identified", True),
        "sso_cc_identified":          tc.get("sso_cc_identified", True),
        "has_transaction_processing": tc.get("has_transaction_processing", True),
        "single_user_entity":         tc.get("single_user_entity", False),
        "has_ai_scope_exclusion":     tc.get("has_ai_scope_exclusion", False),
        "has_other_information":      tc.get("has_other_information", True),
        "addressee_choice":           tc.get("addressee_choice", "Management"),
    }


# ── Helpers ────────────────────────────────────────────────────────────────────

def upload_file(file_bytes, filename, api_base, api_key):
    url = f"{api_base.rstrip('/')}/files/upload"
    resp = requests.post(
        url,
        headers={"Authorization": f"Bearer {api_key}"},
        files={"file": (filename, file_bytes, "application/octet-stream")},
        data={"user": "streamlit-user"},
        timeout=120,
        verify=False,
    )
    resp.raise_for_status()
    return resp.json()["id"]


def to_str(v) -> str:
    """Ensure any workflow output value is a plain string before passing as input."""
    if v is None:
        return ""
    if isinstance(v, str):
        return v
    if isinstance(v, (list, dict)):
        return json.dumps(v, ensure_ascii=False)
    return str(v)


def subservice_org_for_dify(raw) -> str:
    """Reduce Subservice_org to the 'Name | Services' form the Dify backend
    expects before sending it to any workflow.

    The UI text area accepts an optional short-name column
    ('Name | Short Name | Services'), but that short name is consumed only by
    the local EY-template fill (build_substitutions). The backend code node
    parses with split('|', 1) — a 3-column line would put the short name into
    the services field and break the generated bullets/table. Stripping the
    short name here keeps the backend contract unchanged (no Dify re-import).
    """
    if not raw:
        return raw
    out = []
    for line in raw.splitlines():
        s = line.strip()
        if not s:
            continue
        parts = [p.strip() for p in s.split("|")]
        if len(parts) >= 3:
            services = "|".join(parts[2:]).strip()
            out.append(f"{parts[0]} | {services}")
        else:
            out.append(s)
    return "\n".join(out)


def _spinner_html(text: str) -> str:
    """A single in-place status line with an animated CSS spinner. Reused for
    both the node-progress counter and the final build message so the same
    placeholder can be overwritten in place (no second spinner stacked below)."""
    return (
        "<style>"
        "@keyframes _nd_spin{to{transform:rotate(360deg)}}"
        "._nd_s{display:inline-block;width:13px;height:13px;"
        "border:2px solid rgba(180,180,180,0.3);border-top-color:#aaa;"
        "border-radius:50%;animation:_nd_spin 0.75s linear infinite;"
        "vertical-align:middle;margin-right:6px}"
        "</style>"
        f'<div style="font-size:0.95em;padding:3px 0">'
        f'<span class="_nd_s"></span>{text}</div>'
    )


def run_workflow(inputs, api_base, api_key, status_placeholder=None):
    """
    Calls the Dify workflow API in streaming mode to avoid nginx 504 timeouts.
    Parses SSE events and returns the final outputs dict from workflow_finished.
    """
    url = f"{api_base.rstrip('/')}/workflows/run"

    resp = requests.post(
        url,
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json={"inputs": inputs, "response_mode": "streaming", "user": "streamlit-user"},
        stream=True,
        timeout=(30, 1800),  # (connect timeout, read timeout in seconds)
        verify=False,
    )
    resp.raise_for_status()

    outputs = {}
    node_count = 0
    workflow_finished_received = False

    for raw_line in resp.iter_lines():
        if not raw_line:
            continue
        line = raw_line.decode("utf-8") if isinstance(raw_line, bytes) else raw_line
        if not line.startswith("data: "):
            continue
        try:
            event = json.loads(line[6:])
        except json.JSONDecodeError:
            continue

        event_type = event.get("event", "")

        if event_type == "node_finished":
            node_count += 1
            node_data = event.get("data", {})
            node_title = str(node_data.get("title", ""))
            node_status = str(node_data.get("status", ""))
            node_error = str(node_data.get("error", "")) if node_status == "failed" else ""
            # Release the (potentially very large) event dict before any
            # Streamlit UI call — the MAIN-Output end-node payload contains
            # all 19 output fields and can be several hundred KB.
            event = node_data = None
            if status_placeholder:
                status_placeholder.markdown(
                    _spinner_html(
                        f"Nodes completed: {node_count}&nbsp;&nbsp;(last: {node_title})"
                    ),
                    unsafe_allow_html=True,
                )
            if node_status == "failed":
                raise RuntimeError(f"Workflow node failed: {node_error or f'node {node_title!r} (#{node_count})'}")

        elif event_type == "workflow_finished":
            workflow_finished_received = True
            data = event.get("data", {})
            if data.get("status") == "failed":
                raise RuntimeError(f"Workflow failed: {data.get('error', 'unknown error')}")
            outputs = data.get("outputs", {})
            break

        elif event_type == "error":
            raise RuntimeError(event.get("message", "Streaming error from Dify"))

    if not workflow_finished_received:
        raise RuntimeError(
            f"Workflow stream ended unexpectedly after {node_count} node(s) "
            "without a completion event — the workflow may have crashed or timed out on the Dify side."
        )

    return outputs


FONT_LATIN   = "Times New Roman"
FONT_CHINESE = "黑体"


def _apply_fonts(run):
    """Set Times New Roman for Latin characters, 华文楷体 for Chinese characters.
    Word automatically picks the right one per character based on Unicode range."""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"),    FONT_LATIN)
    rFonts.set(qn("w:hAnsi"),    FONT_LATIN)
    rFonts.set(qn("w:eastAsia"), FONT_CHINESE)
    rFonts.set(qn("w:cs"),       FONT_LATIN)


def _set_style_fonts(style):
    """Apply the same dual-font setting (Times New Roman + 华文楷体) at the
    paragraph-style level."""
    rPr = style.element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        style.element.append(rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"),    FONT_LATIN)
    rFonts.set(qn("w:hAnsi"),    FONT_LATIN)
    rFonts.set(qn("w:eastAsia"), FONT_CHINESE)
    rFonts.set(qn("w:cs"),       FONT_LATIN)


def _set_cell_background(cell, fill_hex):
    """Apply a solid background fill to a table cell. fill_hex e.g. 'D9D9D9'."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _add_table_from_md(doc, table_lines):
    """Parse markdown table lines and add a Word table with a gray header row."""
    if len(table_lines) < 2:
        return

    def parse_row(line):
        return [cell.strip() for cell in line.strip().strip("|").split("|")]

    def is_separator(cells):
        return all(re.match(r"^:?-+:?$", c) for c in cells if c)

    headers = parse_row(table_lines[0])
    num_cols = len(headers)

    data_rows = []
    for line in table_lines[1:]:
        cells = parse_row(line)
        if not is_separator(cells):
            data_rows.append(cells)

    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.style = "Table Grid"
    _set_table_borders(table)

    # Header row — gray background, bold text
    for j, header_text in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header_text)
        run.bold = True
        _apply_fonts(run)
        _set_cell_background(cell, "D9D9D9")

    # Data rows — no background
    for i, row_data in enumerate(data_rows):
        for j in range(num_cols):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            text = row_data[j] if j < len(row_data) else ""
            _inline(cell.paragraphs[0], text)


def _add_numpipe_table(doc, lines, language="English"):
    """Parse 'N | Term | Description' rows (no leading pipe) and add a Word table."""
    rows = []
    for line in lines:
        parts = [c.strip() for c in line.strip().split("|") if c.strip()]
        if parts:
            rows.append(parts)
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    if language.startswith("中"):
        hdr = ("编号", "名词/系统名称", "名词解释/系统简介")
    else:
        hdr = ("SN", "Term/Application Name", "Terminology/System Introduction")
    header_row = list(hdr[:num_cols])
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = "Table Grid"
    _set_table_borders(table)
    for j, h in enumerate(header_row):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        _apply_fonts(run)
        _set_cell_background(cell, "D9D9D9")
    for i, row_data in enumerate(rows):
        for j in range(num_cols):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            text = row_data[j] if j < len(row_data) else ""
            run = cell.paragraphs[0].add_run(text)
            _apply_fonts(run)
    _set_col_widths(table, [1.2, 5.0, 9.7])
    _set_repeat_header(table)


def _set_table_borders(tbl):
    """Set explicit solid single borders (outline + inside grid) on a table.

    Generated tables reference the 'Table Grid' style, which is not defined in
    the EY template document the sections are merged into, so the style-based
    borders vanish in the complete report. Borders set on the table itself
    survive the merge."""
    tblPr = tbl._tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tail = next(
            (ch for ch in tblPr if ch.tag in
             (qn("w:tblLayout"), qn("w:tblCellMar"), qn("w:tblLook"))),
            None,
        )
        if tail is not None:
            tail.addprevious(borders)
        else:
            tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")


def _set_col_widths(tbl, widths_cm):
    """Force fixed column widths (list of cm values) on a table."""
    tblPr = tbl._tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in tbl.rows:
        for j, w_cm in enumerate(widths_cm):
            if j < len(row.cells):
                row.cells[j].width = Cm(w_cm)


def _set_repeat_header(table):
    """Make the first row repeat as a header row when the table spans pages."""
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def _add_heading(doc, text, level):
    """Add a Heading-N paragraph with level-specific formatting:
    H1 = bold only | H2 = bold + italic | H3 = italic + underline | H4+ = bold only"""
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    # Explicit outline level: the "Heading N" style reference is lost when the
    # section is merged into the EY template document, which empties Word's
    # navigation pane. outlineLvl keeps the entry without changing formatting.
    _set_outline_level(p, level - 1)
    run = p.add_run(text)
    _apply_fonts(run)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(11)
    if level == 1:
        run.font.bold = True
        run.font.italic = False
    elif level == 2:
        run.font.bold = True
        run.font.italic = True
    elif level == 3:
        run.font.bold = False
        run.font.italic = True
        run.font.underline = True
    else:
        run.font.bold = True
        run.font.italic = False
    blank = doc.add_paragraph()
    blank.paragraph_format.space_after  = Pt(0)
    blank.paragraph_format.space_before = Pt(0)


def _pin_list_numpr(para, doc, left_twips=720, hanging_twips=360):
    """Copy numPr from the paragraph's List Bullet/List Number style onto the
    paragraph itself, and set explicit indent to match the MA template
    (left=1.27 cm, hanging=0.635 cm).

    python-docx stores numPr only in the style definition, not per-paragraph.
    When merged into the EY template docx (which lacks those styles) the bullet
    or number disappears. Inlining numPr makes it survive the merge, and the
    remap/inject machinery in merge_docx_sections will fix the numId so it
    references the injected numbering definitions correctly.
    """
    style_name = para.style.name if para.style else None
    if style_name not in ("List Bullet", "List Number"):
        return
    style = doc.styles[style_name]
    sPPr = style.element.pPr
    if sPPr is None:
        return
    sNumPr = sPPr.find(qn("w:numPr"))
    if sNumPr is None:
        return
    sNumId = sNumPr.find(qn("w:numId"))
    sIlvl  = sNumPr.find(qn("w:ilvl"))
    numId_val = sNumId.get(qn("w:val")) if sNumId is not None else None
    if not numId_val:
        return
    ilvl_val = (sIlvl.get(qn("w:val")) if sIlvl is not None else None) or "0"

    pPr = para._p.get_or_add_pPr()
    existing = pPr.find(qn("w:numPr"))
    if existing is not None:
        pPr.remove(existing)
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), ilvl_val)
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), numId_val)
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)

    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:left"),    str(left_twips))
    ind.set(qn("w:hanging"), str(hanging_twips))


def markdown_to_docx(md_text: str, language: str = "English") -> bytes:
    doc = Document()

    _set_style_fonts(doc.styles["Normal"])
    doc.styles["Normal"].font.size = Pt(11)
    doc.styles["Normal"].paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.styles["Normal"].paragraph_format.space_after  = Pt(0)
    doc.styles["Normal"].paragraph_format.space_before = Pt(0)

    lines = md_text.split("\n")
    i = 0
    # Track whether the last thing added was a blank paragraph so we never
    # emit more than one consecutive blank line regardless of how many '\n'
    # sequences the LLM produced.  Start True so leading blank lines are
    # silently dropped.
    last_was_blank = True

    while i < len(lines):
        line = lines[i]

        # ── Markdown table block ───────────────────────────────────────────
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_table_from_md(doc, table_lines)
            last_was_blank = False
            continue

        # ── Number-pipe table  (N | Term | Description) ────────────────────
        if re.match(r"^\d+\s*\|", line.strip()):
            npt_lines = []
            while i < len(lines) and re.match(r"^\d+\s*\|", lines[i].strip()):
                npt_lines.append(lines[i])
                i += 1
            _add_numpipe_table(doc, npt_lines, language)
            last_was_blank = False
            continue

        # ── Headings ───────────────────────────────────────────────────────
        # H1=bold | H2=bold+italic | H3=italic+underline | H4+=bold
        if line.startswith("#### "):
            _add_heading(doc, line[5:].strip(), 4)
            last_was_blank = True  # _add_heading appends a blank internally

        elif line.startswith("### "):
            _add_heading(doc, line[4:].strip(), 3)
            last_was_blank = True

        elif line.startswith("## "):
            _add_heading(doc, line[3:].strip(), 2)
            last_was_blank = True

        elif line.startswith("# "):
            _add_heading(doc, line[2:].strip(), 1)
            last_was_blank = True

        # ── Bullet lists ───────────────────────────────────────────────────
        elif re.match(r"^[-*+] ", line):
            p = doc.add_paragraph(style="List Bullet")
            _inline_bullet(p, line[2:].strip())
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _pin_list_numpr(p, doc)
            last_was_blank = False

        elif re.match(r"^[•·]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            _inline_bullet(p, re.sub(r"^[•·]\s+", "", line))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _pin_list_numpr(p, doc)
            last_was_blank = False

        elif re.match(r"^\d+\. ", line):
            # Collect the whole consecutive numbered block, keeping the SN
            numbered_items = []
            while i < len(lines) and re.match(r"^\d+\. ", lines[i]):
                m = re.match(r"^(\d+)\. (.*)", lines[i])
                numbered_items.append((m.group(1), m.group(2).strip()))
                i += 1
            # If every item follows "Term: Description", convert to a 3-column table
            _def_pat = re.compile(r"^(?!https?:)([^:]{1,80}):\s+(.*)", re.DOTALL)
            texts = [rest for _, rest in numbered_items]
            if len(texts) >= 2 and all(_def_pat.match(t) for t in texts):
                if language.startswith("中"):
                    hdr = ("编号", "名词/系统名称", "名词解释/系统简介")
                else:
                    hdr = ("SN", "Term/Application Name", "Terminology/System Introduction")
                tbl = doc.add_table(rows=1 + len(numbered_items), cols=3)
                tbl.style = "Table Grid"
                _set_table_borders(tbl)
                for j, h in enumerate(hdr):
                    cell = tbl.rows[0].cells[j]
                    cell.text = ""
                    r = cell.paragraphs[0].add_run(h)
                    r.bold = True
                    _apply_fonts(r)
                    _set_cell_background(cell, "D9D9D9")
                for idx, (sn, rest) in enumerate(numbered_items):
                    m = _def_pat.match(rest)
                    term, desc = m.group(1).strip(), m.group(2)
                    for j, text in enumerate((sn, term, desc)):
                        cell = tbl.rows[idx + 1].cells[j]
                        cell.text = ""
                        run = cell.paragraphs[0].add_run(text)
                        _apply_fonts(run)
                _set_col_widths(tbl, [1.2, 5.0, 9.7])
                _set_repeat_header(tbl)
            else:
                for _, rest in numbered_items:
                    p = doc.add_paragraph(style="List Number")
                    _inline_bullet(p, rest)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    _pin_list_numpr(p, doc)
            last_was_blank = False
            continue  # i already advanced past the block

        # ── Blank line — emit at most one consecutive blank paragraph ──────
        elif line.strip() == "":
            if not last_was_blank:
                blank = doc.add_paragraph()
                blank.paragraph_format.space_after  = Pt(0)
                blank.paragraph_format.space_before = Pt(0)
                last_was_blank = True

        # ── Plain-text heading (no # marker: short line, no terminal punct) ─
        # Excludes both ASCII and Chinese terminal punctuation so that Chinese
        # sentences ending with 。？！，；：are not misidentified as headings.
        elif line.strip() and len(line.strip()) <= 130 and not line.rstrip().endswith(
            ('.', '?', '!', ',', ';', ':', '。', '？', '！', '，', '；', '：')
        ):
            _add_heading(doc, line.strip(), 2)
            last_was_blank = True

        # ── Normal paragraph ───────────────────────────────────────────────
        else:
            p = doc.add_paragraph()
            _inline(p, line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            last_was_blank = False

        i += 1

    # ── Disable "Snap to grid when document grid is defined" on every paragraph ──
    # When a document grid is defined, Word snaps each line to the grid by
    # default (the checkbox is ticked), which distorts line spacing in the
    # generated body. Explicitly add <w:snapToGrid w:val="0"/> to each
    # paragraph so the box is unticked throughout the Dify sections.
    def _disable_snap_to_grid(para):
        pPr = para._p.get_or_add_pPr()
        snap = pPr.find(qn("w:snapToGrid"))
        if snap is None:
            snap = OxmlElement("w:snapToGrid")
            # Per the CT_PPr schema, snapToGrid precedes these elements.
            anchor = None
            for _tag in ("w:spacing", "w:ind", "w:jc", "w:rPr", "w:sectPr", "w:pPrChange"):
                anchor = pPr.find(qn(_tag))
                if anchor is not None:
                    break
            if anchor is not None:
                anchor.addprevious(snap)
            else:
                pPr.append(snap)
        snap.set(qn("w:val"), "0")

    for para in doc.paragraphs:
        _disable_snap_to_grid(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _disable_snap_to_grid(para)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _inline(paragraph, text):
    for part in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text):
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            _apply_fonts(run)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            _apply_fonts(run)
        else:
            run = paragraph.add_run(part)
            _apply_fonts(run)


def _inline_bullet(paragraph, text):
    """For bullet items: if markdown bold markers are present use _inline as-is;
    otherwise auto-bold the 'Header: content' pattern before the first colon."""
    if "**" in text:
        _inline(paragraph, text)
        return
    # Match 'Header: content' — colon within 80 chars, not a URL scheme (http/https)
    m = re.match(r"^(?!https?:)([^:]{1,80}):\s*(.*)", text, re.DOTALL)
    if m:
        header, content = m.group(1).strip(), m.group(2)
        run = paragraph.add_run(header + ": ")
        run.bold = True
        _apply_fonts(run)
        if content:
            run2 = paragraph.add_run(content)
            _apply_fonts(run2)
    else:
        _inline(paragraph, text)


def build_final_document(result_text, ui, tc):
    """Build the final .docx bytes (+ filename) from the SUB2 markdown result.

    Returns (bytes, filename). In complete mode the EY MA/AR templates are filled
    and merged with the Dify sections; otherwise the Dify sections are rendered
    on their own. Any failure in the complete path falls back to Dify-only.
    Caller is responsible for wrapping this in a spinner if desired.
    """
    dify_bytes = markdown_to_docx(result_text, ui.get("Output_language", "English"))

    if (tc.get("generate_complete")
            and tc.get("ar_template_path")
            and tc.get("ma_template_path")):
        subs  = build_substitutions(ui, tc)
        flags = build_flags(tc)
        _lang = ui.get("Output_language", "English")
        _lh_path = tc.get("letterhead_path")
        _use_lh  = bool(_lh_path) and os.path.isfile(_lh_path)
        try:
            ma_bytes = fill_and_process_template(tc["ma_template_path"], subs, flags, _lang)
            ar_bytes = fill_and_process_template(tc["ar_template_path"], subs, flags, _lang)
            # Order is MA, AR, Dify → the AR section is the 2nd (index 1).
            built = enforce_line_spacing(
                merge_docx_sections(ma_bytes, ar_bytes, dify_bytes, split_sections=_use_lh)
            )
            built = strip_page_top_empty_paragraphs(built)
            if _use_lh:
                built = inject_ar_letterhead(built, _lh_path, ar_index=1)
            fname = (
                f"{ui.get('Co_short_name', 'Report')}_"
                f"{ui.get('Report_type', '').replace(' ', '_')}_Complete_Report.docx"
            )
            return built, fname
        except Exception as exc:
            st.error(f"Failed to generate complete report: {exc}\n\nFalling back to Dify sections only.")

    built = strip_page_top_empty_paragraphs(enforce_line_spacing(dify_bytes))
    fname = (
        f"{ui.get('Co_short_name', 'Report')}_"
        f"{ui.get('Report_type', '').replace(' ', '_')}_Report.docx"
    )
    return built, fname


# ══════════════════════════════════════════════════════════════════════════════
# STEP 1 — Report Parameters & MAIN Workflow
# ══════════════════════════════════════════════════════════════════════════════
# The input form is shown only until a report exists. Once a report is generated
# (`final_done`) the whole form is hidden and the finished report renders in the
# FINAL RESULT block below; the user clicks Reset to start a new report.
if not final_done:

    st.subheader("Upload Control Matrix File(s)")
    st.caption(
        "1. 请上传Excel大表，其中必须包含Control Matrix sheet；"
        # "2. 若有，上传子服务机构及其服务内容Sheet；"
        # "3. 若有，上传CUEC内容清单Sheet；"
        # "3. 若有，上传子服务机构补充控制清单Sheet；"
        "2. 若有，上传名词解释Sheet；"
        "3. 若适用，上传控制目标Sheet"
    )
    uploaded_files = st.file_uploader(
        "Upload files (Excel, PDF, Word, etc.)",
        accept_multiple_files=True,
    )

    # ── Required fields ───────────────────────────────────────────────────────
    st.subheader("Required Fields")
    req1, req2 = st.columns(2)

    with req1:
        company_name        = st.text_input("Company Name",         max_chars=256)
        co_short_name       = st.text_input("Company Short Name",   max_chars=48)
        system_name         = st.text_input("Service / System Name",max_chars=256)
        period_start    = st.text_input("Report Period Start (as of date if Type1)", placeholder="e.g. 2025-01-01")
        scope_of_report = st.selectbox("Subservice Organization Testing Strategy",
            ["None", "All carve out", "Inclusive"], key="form_scope_of_report")
        industry = st.selectbox("Industry",
                                ["HR", "IaaS", "AI", "SaaS", "Others"])

    with req2:
        report_type     = st.selectbox("Report Type",
            ["SOC1 TYPE1", "SOC1 TYPE2", "SOC2 TYPE1", "SOC2 TYPE2"], key="form_report_type")
        output_language = st.selectbox("Output Language",
            ["English", "中文"], key="form_output_language")
        service_description = st.text_input("Service Description",  max_chars=256)
        period_end = st.text_input("Report Period End (N/A for Type1)", placeholder="e.g. 2025-12-31")
        subservice_org = st.text_area(
                            "Subservice Organization (N/A if no subservice organization)",
                            placeholder="Alibaba Cloud | Alibaba Cloud | Elastic Cloud, Object Storage\nTencent Cloud | Tencent | Cloud Virtual Machine, TencentDB",
                            help=(
                                "One entry per line. Format: Full Name | Short Name | Services\n"
                                "Example: Alibaba Cloud | Alibaba Cloud | Elastic Cloud, Object Storage\n"
                                "If Short Name is omitted (2 columns), Full Name is used as Short Name."
                            ),
                            height=100,
                        )
        if len(subservice_org) > 256:
            st.warning("⚠️ Subservice Organization exceeds 256 characters. Please shorten it.")

    # ── Complete report option ─────────────────────────────────────────────────
    # Placed after Required Fields so the MA/AR template settings below react to
    # the Report Type / Subservice Testing Strategy / Output Language chosen above.
    st.markdown("---")
    generate_complete = st.checkbox(
        "Generate complete report (MA + AR + main sections)",
        value=True,
        help="When checked, the download will include Section I (Management Assertion) and Section II (Independent Auditor's Report) generated from EY templates, followed by the Dify-generated sections. Uncheck to generate Sections III–IV only (existing behaviour).",
    )

    if generate_complete:
        with st.expander("Complete Report Settings", expanded=True):
            cr1, cr2 = st.columns(2)

            with cr1:
                _std_options = get_standard_options(report_type)
                standard = st.selectbox("Standard", _std_options, key="cr_standard")
                report_date  = st.text_input(
                    "Report Signing Date (YYYY-MM-DD)",
                    placeholder="e.g. 2026-01-30",
                    key="cr_report_date",
                    help="Formatted automatically: \"January 30, 2026\" in English reports, "
                         "\"2026年1月30日\" in Chinese reports. "
                         "Other text is inserted into the report as-is.",
                )
                signing_city = st.text_input("Signing City", placeholder="e.g. Shanghai/上海", key="cr_signing_city")

                _lh_files = list_letterheads()
                _lh_options = ["(None — no letterhead)"] + _lh_files
                letterhead_choice = st.selectbox(
                    "AR Letterhead (EY office)",
                    _lh_options,
                    index=0,
                    key="cr_letterhead",
                    help="EY office letterhead applied to the Auditor's Report (Section II) "
                         "pages only — full letterhead on the first page, continuation header "
                         "after. Files are read from AR_template/letterheads/ (next to the .exe "
                         "when packaged); group + centre the header in Word before saving each "
                         "file. Choose (None) to omit the letterhead.",
                )
                letterhead_path = (
                    os.path.join(LETTERHEAD_DIR, letterhead_choice)
                    if letterhead_choice in _lh_files else None
                )
                if not _lh_files:
                    st.caption(
                        "ℹ️ No letterhead files found in AR_template/letterheads/. "
                        "Add EY office letterhead .docx files there to enable the letterhead."
                    )

            with cr2:
                addressee_choice = st.radio(
                    "AR Addressee",
                    ["Management", "Board of Directors"],
                    index=0,
                    key="cr_addressee",
                    help="Controls whether the AR opens 'To the Management of' or 'To the Board of Directors of' followed by the service organization name.",
                )
                cuec_choice = st.radio(
                    "Complementary User Entity Controls (CUEC)",
                    ["Identified", "Not Identified"],
                    index=0,
                    key="cr_cuec",
                )
                has_transaction_processing = st.checkbox(
                    "Includes transaction processing wording",
                    value=True,
                    key="cr_transaction",
                )
                if not has_transaction_processing:
                    st.caption(
                        "⚠️ When unchecked, the 'Systems Function' field "
                        "(Optional Fields section) is used to describe the system's "
                        "function in place of transaction-processing wording."
                    )
                single_user_entity = st.checkbox(
                    "Single user entity report",
                    value=False,
                    key="cr_single_user",
                )
                has_ai_scope_exclusion = st.checkbox(
                    "Subject matter includes AI technology (audit scope excludes AI-specific functions)",
                    value=False,
                    key="cr_ai_scope",
                    help="When checked, includes the paragraph disclosing that AI technology is used in the subject matter but is not within the audit scope. Leave unchecked if the subject matter does not involve AI technology.",
                )
                has_other_information = st.checkbox(
                    "Report includes 'Other Information' section",
                    value=True,
                    key="cr_other_info",
                    help="When unchecked, template paragraphs that refer to the Other Information section (注：如果没有Other Information这一章节，删除本段) are removed.",
                )
                # SSO CC — only shown when SSO != None
                if scope_of_report != "None":
                    sso_cc_choice = st.radio(
                        "SSO Complementary Controls",
                        ["Identified", "Not Identified"],
                        index=0,
                        key="cr_sso_cc",
                    )
                else:
                    sso_cc_choice = "Identified"

            # Template resolution preview (computed every render)
            st.markdown("---")
            _ar_wp, _ar_path = resolve_template(report_type, standard, scope_of_report, output_language, "AR")
            _ma_wp, _ma_path = resolve_template(report_type, standard, scope_of_report, output_language, "MA")

            def _show_template_status(label, wp, path, dir_name):
                if wp is None and isinstance(path, str):
                    # path carries the error message when wp is None
                    st.error(f"{label}: {path}")
                elif wp is None:
                    st.warning(f"{label}: No matching template found for this combination.")
                elif path and os.path.isfile(path):
                    st.info(f"{label}: WP No. {wp} → {os.path.basename(path)}")
                elif isinstance(path, str) and path.startswith("Cannot"):
                    st.error(f"{label}: WP No. {wp} — {path}")
                else:
                    st.warning(f"{label}: WP No. {wp} listed but .docx not found in {dir_name}/")

            _show_template_status("AR template", _ar_wp, _ar_path, "AR_template")
            _show_template_status("MA template", _ma_wp, _ma_path, "MA_template")

            # Where the MA/AR templates are being read from this session.
            if TEMPLATE_DIR_WARNING:
                st.warning(f"⚠️ {TEMPLATE_DIR_WARNING}")
            elif TEMPLATE_DIR_SOURCE == "sharepoint":
                st.caption(f"📂 Templates source: SharePoint library ({SP_SITE_URL}) — fetched this session")
            elif TEMPLATE_DIR_SOURCE == "feishu":
                st.caption("📂 Templates source: Feishu Drive — fetched this session")
            elif TEMPLATE_DIR_SOURCE == "onedrive":
                st.caption(f"📂 Templates source: synced SharePoint library ({TEMPLATE_BASE_PATH})")
            else:
                st.caption("📂 Templates source: bundled (set TEMPLATE_SOURCE=sharepoint to fetch the latest from SharePoint)")

    else:
        standard                  = ""
        report_date               = ""
        signing_city              = ""
        cuec_choice               = "Identified"
        sso_cc_choice             = "Identified"
        has_transaction_processing = True
        single_user_entity         = False
        has_ai_scope_exclusion     = False
        has_other_information      = True
        addressee_choice           = "Management"
        letterhead_path            = None
        _ar_path                   = None
        _ma_path                   = None



    # ── Optional fields ────────────────────────────────────────────────────────
    st.markdown("---")
    st.subheader("Optional Fields")
    opt1, opt2 = st.columns(2)

    with opt1:
        domain         = st.text_input("Control Domain",                  max_chars=256)
        co_website = st.text_input("Company Website", max_chars=256)
    with opt2:
        systems_function = st.text_input("Systems Function",
                        placeholder="e.g. workflow approval, code management, cloud resource management",
                        help="Optional. Describe the purpose of the internal supporting systems listed above.",
                        max_chars=256)
        system_extra = st.text_input("Internal Supporting Systems",
                         placeholder="e.g. Feishu Platform, Gitlab Platform, Alibaba Cloud Console",
                         help="Optional. List the internal systems used to support operations. If left blank, the workflow will auto-extract from the Control Matrix.",
                         max_chars=256)

    st.subheader("Trust Service Criteria (SOC2 only)")
    tsc_cols = st.columns(5)
    is_security              = tsc_cols[0].checkbox("Security",             key="form_tsc_security")
    is_availability          = tsc_cols[1].checkbox("Availability",         key="form_tsc_availability")
    is_processing_integrity  = tsc_cols[2].checkbox("Processing Integrity", key="form_tsc_processing")
    is_confidentiality       = tsc_cols[3].checkbox("Confidentiality",      key="form_tsc_confidentiality")
    is_privacy               = tsc_cols[4].checkbox("Privacy",              key="form_tsc_privacy")

    st.subheader("User Entity Section")
    # Defaults follow the report type: CUEC for SOC1, UER for SOC2. The report_type
    # is baked into the widget key so the default re-applies when the report type
    # changes, while still letting the user override within a given type.
    ue_cols = st.columns(2)
    # Streamlit drops a widget's key from session_state whenever the widget isn't
    # rendered on a run — including the Reset rerun (the form isn't rendered before
    # st.rerun() fires) — so a plain keyed checkbox re-seeds to the report-type
    # default on a MA+AR-only Reset while every other field is kept. We keep a
    # shadow copy in a plain (non-widget) key the GC can't touch: seed it with the
    # report-type default, mirror every toggle into it via on_change, and seed the
    # checkbox from it through `value=`. The selection then survives a MA+AR-only /
    # mid-run Reset. After a COMPLETE report the form is hidden and everything else
    # resets, so the Reset handler clears these shadows in that case only (sidebar).
    #
    # IMPORTANT: we never write the widget key ourselves (`st.session_state[_cuec_key]
    # = …`). Doing that alongside on_change made clicking UER reset CUEC. `value=`
    # is only applied on first render (and after a GC) and is ignored once the key
    # exists, so a normal rerun keeps each box's own state. The key embeds
    # report_type so switching types applies that type's default.
    _cuec_key = f"form_is_cuec_{report_type}"
    _uer_key  = f"form_is_uer_{report_type}"
    _cuec_pref = "pref_" + _cuec_key
    _uer_pref  = "pref_" + _uer_key
    if _cuec_pref not in st.session_state:
        st.session_state[_cuec_pref] = report_type.startswith("SOC1")
    if _uer_pref not in st.session_state:
        st.session_state[_uer_pref] = report_type.startswith("SOC2")

    def _mirror_pref(widget_key):
        st.session_state["pref_" + widget_key] = st.session_state[widget_key]

    # Labels are kept short (acronym only) so they stay on a single line within the
    # half-width column — otherwise the long CUEC label wraps and its help "?" icon
    # drops to the second line, misaligning it with UER's. Full names live in `help`.
    is_cuec = ue_cols[0].checkbox(
        "Include CUEC",
        value=st.session_state[_cuec_pref],
        key=_cuec_key,
        on_change=_mirror_pref, args=(_cuec_key,),
        help="Complementary User Entity Controls — default on for SOC1 reports. "
             "Generated from the control matrix.")
    is_uer = ue_cols[1].checkbox(
        "Include UER",
        value=st.session_state[_uer_pref],
        key=_uer_key,
        on_change=_mirror_pref, args=(_uer_key,),
        help="User Entity Responsibilities — default on for SOC2 reports. "
             "Generated from the control matrix.")

    st.markdown("---")

    run_main = st.button(
        "▶ Run All Steps (1 → 2 → 3)", type="primary",
        use_container_width=True,
    )
    run_ma_ar_only = False
    if generate_complete:
        run_ma_ar_only = st.button(
            "🧪 Generate MA + AR only (templates, no Dify)",
            use_container_width=True,
            help="Fills the Section I + II templates using the fields above "
                 "without running the Dify workflow.",
        )

    # Live Run-All status placeholders — the "⏳ Step N — Running…" banner and the
    # "Nodes completed…" line. Created at a fixed position every run so they stay
    # empty (invisible) unless a Dify run writes to them, and so the MA+AR-only
    # path can explicitly clear any banner left over from an interrupted Run-All.
    step_label  = st.empty()
    node_status = st.empty()
    ma_ar_dl    = st.empty()   # holds the "⬇ Download MA + AR" button

    # A click on Run All Steps or Generate MA + AR only starts new (blocking)
    # processing — immediately clear any leftover Run-All banner and stale MA+AR
    # download button so they don't linger on screen while it runs.
    if run_main or run_ma_ar_only:
        step_label.empty()
        node_status.empty()
        ma_ar_dl.empty()

    # ── MA + AR only: fill the templates from the fields above, no Dify run ────
    if generate_complete:
        if run_ma_ar_only:
            _ar_wp_t, _ar_path_t = resolve_template(report_type, standard, scope_of_report, output_language, "AR")
            _ma_wp_t, _ma_path_t = resolve_template(report_type, standard, scope_of_report, output_language, "MA")
            # Require the same form fields as "Run All Steps", minus the Dify-only
            # inputs (file upload and API key are not needed to fill the templates).
            ma_ar_errors = []
            if not company_name:        ma_ar_errors.append("Company Name is required.")
            if not co_short_name:       ma_ar_errors.append("Company Short Name is required.")
            if not system_name:         ma_ar_errors.append("Service/System Name is required.")
            if not service_description: ma_ar_errors.append("Service Description is required.")
            if not period_start:        ma_ar_errors.append("Report Period Start is required.")
            if report_type.endswith("TYPE2") and not period_end:
                ma_ar_errors.append("Report Period End is required for Type 2 reports.")
            if report_type.startswith("SOC2") and not any([
                is_security, is_availability, is_processing_integrity,
                is_confidentiality, is_privacy,
            ]):
                ma_ar_errors.append("At least one Trust Service Criteria must be selected for SOC2 reports.")
            if len(subservice_org) > 256:
                ma_ar_errors.append("Subservice Organization exceeds 256 characters")
            if (scope_of_report != "None") and (not subservice_org):
                ma_ar_errors.append("Please input Subservice Organizations and its service provided")
            if not report_date:
                ma_ar_errors.append("Report Signing Date is required when generating a complete report.")
            if not signing_city:
                ma_ar_errors.append("Signing City is required when generating a complete report.")
            if not has_transaction_processing and not systems_function:
                ma_ar_errors.append(
                    "Systems Function is required when 'Includes transaction processing wording' "
                    "is unchecked — please fill in the Systems Function field."
                )
            if not (_ar_path_t and os.path.isfile(_ar_path_t)):
                ma_ar_errors.append(f"AR template not available: {_ar_path_t or 'no matching template found'}")
            if not (_ma_path_t and os.path.isfile(_ma_path_t)):
                ma_ar_errors.append(f"MA template not available: {_ma_path_t or 'no matching template found'}")
            if ma_ar_errors:
                for _e in ma_ar_errors:
                    st.error(_e)
            else:
                _test_ui = {
                    "Company_name":           company_name,
                    "Co_short_name":          co_short_name,
                    "System_or_service_name": system_name,
                    "Service_description":    service_description,
                    "Period_start":           period_start,
                    "Period_end":             period_end,
                    "Report_type":            report_type,
                    "Output_language":        output_language,
                    "Subservice_org":         subservice_org,
                    "Systems_function":       systems_function    or "",
                    "is_Security":             is_security,
                    "is_Availability":         is_availability,
                    "is_Processing_Integrity": is_processing_integrity,
                    "is_Confidentiality":      is_confidentiality,
                    "is_Privacy":              is_privacy,
                }
                _test_tc = {
                    "report_date":                report_date,
                    "signing_city":               signing_city,
                    "cuec_identified":            cuec_choice == "Identified",
                    "sso_cc_identified":          sso_cc_choice == "Identified",
                    "has_transaction_processing": has_transaction_processing,
                    "single_user_entity":         single_user_entity,
                    "has_ai_scope_exclusion":     has_ai_scope_exclusion,
                    "has_other_information":      has_other_information,
                    "addressee_choice":           addressee_choice,
                }
                try:
                    with st.spinner("Generating MA + AR sections…"):
                        _t_subs  = build_substitutions(_test_ui, _test_tc)
                        _t_flags = build_flags(_test_tc)
                        _t_ma    = fill_and_process_template(_ma_path_t, _t_subs, _t_flags, output_language)
                        _t_ar    = fill_and_process_template(_ar_path_t, _t_subs, _t_flags, output_language)
                        _t_use_lh = bool(letterhead_path) and os.path.isfile(letterhead_path)
                        # Order is MA, AR → the AR section is the 2nd (index 1).
                        _t_merged = enforce_line_spacing(
                            merge_docx_sections(_t_ma, _t_ar, split_sections=_t_use_lh)
                        )
                        _t_merged = strip_page_top_empty_paragraphs(_t_merged)
                        if _t_use_lh:
                            _t_merged = inject_ar_letterhead(_t_merged, letterhead_path, ar_index=1)
                    st.session_state["ma_ar_only"] = {
                        "bytes": _t_merged,
                        "filename": (
                            f"{(co_short_name or 'Test')}_{report_type.replace(' ', '_')}"
                            f"_MA_AR.docx"
                        ),
                    }
                except Exception as _exc:
                    st.session_state.pop("ma_ar_only", None)
                    st.error(f"MA + AR generation failed: {_exc}")
                    with st.expander("Show error details (traceback)", expanded=True):
                        st.code(traceback.format_exc())

        # Clicking "Run All Steps" starts a Dify run; hide any pending MA+AR
        # download so its stale data can't be re-downloaded mid-workflow (a second
        # click on the download button interrupts the running workflow).
        if run_main:
            st.session_state.pop("ma_ar_only", None)

        # Render the download into its fixed placeholder so it sits in a stable
        # slot that the clear-block above can wipe the instant a new run starts.
        # The button persists across downloads — it only disappears when a new
        # "Run All Steps" / "Generate MA + AR only" click starts fresh processing.
        if st.session_state.get("ma_ar_only"):
            ma_ar_dl.download_button(
                label="⬇ Download MA + AR (.docx)",
                data=st.session_state["ma_ar_only"]["bytes"],
                file_name=st.session_state["ma_ar_only"]["filename"],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    if run_main:
        errors = []
        if not key_main:   errors.append("MAIN Workflow API key is not set. Add DIFY_API_KEY_MAIN to the .env file next to the executable.")
        if not company_name: errors.append("Company Name is required.")
        if not co_short_name: errors.append("Company Short Name is required.")
        if not system_name: errors.append("Service/System Name is required.")
        if not service_description: errors.append("Service Description is required.")
        if not period_start: errors.append("Report Period Start is required.")
        if report_type.endswith("TYPE2") and not period_end:
            errors.append("Report Period End is required for Type 2 reports.")
        if report_type.startswith("SOC2") and not any([
            is_security, is_availability, is_processing_integrity,
            is_confidentiality, is_privacy,
        ]):
            errors.append("At least one Trust Service Criteria must be selected for SOC2 reports.")
        if not uploaded_files: errors.append("At least one file must be uploaded.")
        if len(subservice_org) > 256:
            errors.append("Subservice Organization exceeds 256 characters")
        if (scope_of_report != "None") and (not subservice_org):
            errors.append("Please input Subservice Organizations and its service provided")

        if generate_complete:
            if not report_date:
                errors.append("Report Signing Date is required when generating a complete report.")
            if not signing_city:
                errors.append("Signing City is required when generating a complete report.")
            if not has_transaction_processing and not systems_function:
                errors.append(
                    "Systems Function is required when 'Includes transaction processing wording' "
                    "is unchecked — please fill in the Systems Function field."
                )

        if errors:
            for e in errors:
                st.error(e)
            st.stop()

        # Re-resolve templates using the current form's report_type, sso and language
        if generate_complete:
            _ar_wp_final, _ar_path_final = resolve_template(report_type, standard, scope_of_report, output_language, "AR")
            _ma_wp_final, _ma_path_final = resolve_template(report_type, standard, scope_of_report, output_language, "MA")
            # Treat error strings (non-file paths) as missing
            if _ar_path_final and not os.path.isfile(_ar_path_final):
                st.warning(f"AR template issue: {_ar_path_final} — complete report will omit Section II.")
                _ar_path_final = None
            elif not _ar_path_final:
                st.warning("AR template file not found for this combination — complete report will omit Section II.")
            if _ma_path_final and not os.path.isfile(_ma_path_final):
                st.warning(f"MA template issue: {_ma_path_final} — complete report will omit Section I.")
                _ma_path_final = None
            elif not _ma_path_final:
                st.warning("MA template file not found for this combination — complete report will omit Section I.")
        else:
            _ar_path_final = None
            _ma_path_final = None

        # Upload files
        with st.spinner("Uploading file(s) to Dify…"):
            try:
                file_ids = []
                for uf in uploaded_files:
                    fid = upload_file(uf.read(), uf.name, api_base, key_main)
                    file_ids.append({
                        "transfer_method": "local_file",
                        "upload_file_id": fid,
                        "type": "document",
                    })
            except requests.HTTPError as e:
                st.error(f"File upload failed: {e.response.status_code} — {e.response.text}")
                st.stop()
            except Exception as e:
                st.error(f"File upload error: {e}")
                st.stop()

        # EN reports: fix names typed in all-lowercase / ALL-CAPS so the Dify
        # sections use the same normalised names as the MA/AR templates.
        if output_language != "中文":
            company_name  = _capitalize_name(company_name)
            co_short_name = _capitalize_name(co_short_name)
            system_name   = _capitalize_name(system_name)

        # Store user inputs for later steps
        st.session_state["user_inputs"] = {
            "Report_type": report_type,
            "Output_language": output_language,
            "Company_name": company_name,
            "Co_short_name": co_short_name,
            "Industry": industry,
            "System_or_service_name": system_name,
            "Service_description": service_description,
            "System": system_extra,
            "Systems_function": systems_function,
            "Domain": domain,
            "Period_start": period_start,
            "Period_end": period_end,
            "Subservice_org": subservice_org,
            "Scope_of_the_report": scope_of_report,
            "Co_website": co_website,
            "is_Security":             is_security,
            "is_Availability":         is_availability,
            "is_Processing_Integrity": is_processing_integrity,
            "is_Confidentiality":      is_confidentiality,
            "is_Privacy":              is_privacy,
            "is_CUEC":                 is_cuec,
            "is_UER":                  is_uer,
        }

        # Store template config for download step
        st.session_state["template_config"] = {
            "generate_complete":          generate_complete,
            "standard":                   standard,
            "report_date":                report_date,
            "signing_city":               signing_city,
            "cuec_identified":            cuec_choice == "Identified",
            "sso_cc_identified":          sso_cc_choice == "Identified",
            "has_transaction_processing": has_transaction_processing,
            "single_user_entity":         single_user_entity,
            "has_ai_scope_exclusion":     has_ai_scope_exclusion,
            "has_other_information":      has_other_information,
            "addressee_choice":           addressee_choice,
            "ar_template_path":           _ar_path_final,
            "ma_template_path":           _ma_path_final,
            "letterhead_path":            letterhead_path,
        }

        inputs_main = {
            **st.session_state["user_inputs"],
            "Subservice_org": subservice_org_for_dify(subservice_org),
            "File_input": file_ids,
        }

        # step_label / node_status are the shared placeholders created above.
        if True:
            # ── Step 1 ────────────────────────────────────────────────────────
            step_label.info("⏳ Step 1 — Running MAIN workflow — this may take several minutes…")
            try:
                outputs_main = run_workflow(inputs_main, api_base, key_main, node_status)
            except requests.HTTPError as e:
                st.error(f"MAIN workflow error: {e.response.status_code} — {e.response.text}")
                st.stop()
            except Exception as e:
                st.error(f"MAIN workflow error: {e}")
                st.stop()
            if outputs_main.get("Error") or outputs_main.get("Error_ORG"):
                st.error(f"MAIN workflow returned an error:\n{outputs_main.get('Error') or outputs_main.get('Error_ORG')}")
                st.stop()
            st.session_state["main_outputs"] = outputs_main
            status_bar.markdown(_status_html("✅", "🟡", "⚪"), unsafe_allow_html=True)

            # ── Step 2 ────────────────────────────────────────────────────────
            if not key_sub1:
                st.error("SUB1 API key is required (set in sidebar).")
                st.stop()
            step_label.info("⏳ Step 2 — Running SUB1 workflow — this may take several minutes…")
            mo = outputs_main
            ui = st.session_state["user_inputs"]
            inputs_sub1 = {
                "overview_section":               to_str(mo.get("overview_section")),
                "principals_section":             to_str(mo.get("principals_section")),
                "scope_section":                  to_str(mo.get("scope_section")),
                "org_overview_section":           to_str(mo.get("org_overview_section")),
                "clean_pairs":                    to_str(mo.get("clean_pairs")),
                "activity_domains_text":          to_str(mo.get("activity_domains_text")),
                "entity_domain_controls":         to_str(mo.get("entity_domain_controls")),
                "needing_assignment":             to_str(mo.get("needing_assignment")),
                "CO_list_text":                   to_str(mo.get("CO_list_text")),
                "CUEC_json":                      to_str(mo.get("CUEC_json")),
                "CSOC_json":                      to_str(mo.get("CSOC_json")),
                "Terminology_json":               to_str(mo.get("Terminology_json")),
                "control_list_text":              to_str(mo.get("control_list_text")),
                "org_structure":                  to_str(mo.get("org_structure")),
                "website_content":                to_str(mo.get("website_content")),
                "activity_control_objectives_json": to_str(mo.get("activity_control_objectives_json")),
                "entity_control_objective":       to_str(mo.get("entity_control_objective")),
                "entity_domain_packs_direct":     to_str(mo.get("entity_domain_packs_direct")),
                "rag_context":                    to_str(mo.get("rag_context")),
                "Report_type":            ui.get("Report_type", ""),
                "Output_language":        ui.get("Output_language", ""),
                "Company_name":           ui.get("Company_name", ""),
                "Co_short_name":          ui.get("Co_short_name", ""),
                "Industry":               ui.get("Industry", ""),
                "System_or_service_name": ui.get("System_or_service_name", ""),
                "Subservice_org":         subservice_org_for_dify(ui.get("Subservice_org", "")),
                "Scope_of_the_report":    ui.get("Scope_of_the_report", ""),
                "Domain":                 ui.get("Domain", ""),
                "Period_start":           ui.get("Period_start", ""),
                "Period_end":             ui.get("Period_end", ""),
            }
            try:
                outputs_sub1 = run_workflow(inputs_sub1, api_base, key_sub1, node_status)
            except requests.HTTPError as e:
                st.error(f"SUB1 workflow error: {e.response.status_code} — {e.response.text}")
                st.stop()
            except Exception as e:
                st.error(f"SUB1 workflow error: {e}")
                st.stop()
            st.session_state["sub1_outputs"] = outputs_sub1
            status_bar.markdown(_status_html("✅", "✅", "🟡"), unsafe_allow_html=True)

            # ── Step 3 ────────────────────────────────────────────────────────
            if not key_sub2:
                st.error("SUB2 API key is required (set in sidebar).")
                st.stop()
            step_label.info("⏳ Step 3 — Running SUB2 workflow — this may take several minutes…")
            so = outputs_sub1
            mo = outputs_main
            ui = st.session_state.get("user_inputs", {})
            inputs_sub2 = {
                "overview_section":               to_str(so.get("overview_section")),
                "principals_section":             to_str(so.get("principals_section")),
                "scope_section":                  to_str(so.get("scope_section")),
                "org_overview_section":           to_str(so.get("org_overview_section")),
                "entity_level_section":           to_str(so.get("entity_level_section")),
                "clean_pairs":                    to_str(so.get("clean_pairs")),
                "activity_domains_text":          to_str(so.get("activity_domains_text")),
                "CO_list_text":                   to_str(so.get("CO_list_text")),
                "CUEC_json":                      to_str(so.get("CUEC_json")),
                "CSOC_json":                      to_str(so.get("CSOC_json")),
                "Terminology_json":               to_str(so.get("Terminology_json")),
                "website_content":                to_str(so.get("website_content")),
                "activity_control_objectives_json": to_str(so.get("activity_control_objectives_json")),
                "rag_context":                    to_str(so.get("rag_context")),
                "Subservice_org":         subservice_org_for_dify(to_str(so.get("Subservice_org") or ui.get("Subservice_org"))),
                "Scope_of_the_report":    to_str(so.get("Scope_of_the_report") or ui.get("Scope_of_the_report")),
                "Period_start":           to_str(so.get("Period_start") or ui.get("Period_start")),
                "Period_end":             to_str(so.get("Period_end") or ui.get("Period_end")),
                "Report_type":            to_str(so.get("Report_type") or ui.get("Report_type")),
                "Output_language":        to_str(so.get("Output_language") or ui.get("Output_language")),
                "Company_name":           to_str(so.get("Company_name") or ui.get("Company_name")),
                "Co_short_name":          to_str(so.get("Co_short_name") or ui.get("Co_short_name")),
                "System_or_service_name": to_str(so.get("System_or_service_name") or ui.get("System_or_service_name")),
                "cuec_preformatted":      to_str(mo.get("cuec_preformatted")),
                "UER_json":               to_str(mo.get("UER_json")),
                "control_list_text":      to_str(mo.get("control_list_text")),
                "is_CUEC":                bool(ui.get("is_CUEC", False)),
                "is_UER":                 bool(ui.get("is_UER", False)),
            }
            try:
                outputs_sub2 = run_workflow(inputs_sub2, api_base, key_sub2, node_status)
            except requests.HTTPError as e:
                st.error(f"SUB2 workflow error: {e.response.status_code} — {e.response.text}")
                st.stop()
            except Exception as e:
                st.error(f"SUB2 workflow error: {e}")
                st.stop()
            result = outputs_sub2.get("Result", "")
            if not result:
                st.warning("SUB2 completed but returned no output.")
                st.stop()
            st.session_state["final_result"] = result

            # Build the final .docx now — while we are still rendered inside the
            # Steps container, so the status is visible where the user is already
            # looking instead of off-screen after the rerun. Cached so the
            # post-rerun result section reuses it instead of rebuilding.
            #
            # Overwrite the node-counter line *in place* (same placeholder) with
            # a single stable build message — this removes the "Nodes completed…"
            # spinner instead of stacking a second spinner below it.
            step_label.empty()
            node_status.markdown(
                _spinner_html("Generating and merging MA &amp; AR section (Section I &amp; II)…"),
                unsafe_allow_html=True,
            )
            _built, _fname = build_final_document(
                result, st.session_state.get("user_inputs", {}),
                st.session_state.get("template_config", {}),
            )
            st.session_state["final_bytes"]    = _built
            st.session_state["final_filename"] = _fname

        status_bar.markdown(_status_html("✅", "✅", "✅"), unsafe_allow_html=True)
        st.rerun()



# ══════════════════════════════════════════════════════════════════════════════
# FINAL RESULT — Preview + Download
# ══════════════════════════════════════════════════════════════════════════════
if final_done:
    st.markdown("---")
    st.success("🎉 Report is ready!")

    result_text = st.session_state["final_result"]
    ui  = st.session_state.get("user_inputs", {})
    tc  = st.session_state.get("template_config", {})

    # The document is normally built right after Step 3 (inside the Steps
    # container, so the spinner is visible there). This guarded block is only a
    # fallback for sessions where final_result exists but final_bytes does not
    # (e.g. a restored session). Cached so reruns don't rebuild it.
    if "final_bytes" not in st.session_state:
        with st.spinner("Generating and merging MA & AR section (Section I & II)…"):
            _built, _fname = build_final_document(result_text, ui, tc)
        st.session_state["final_bytes"]    = _built
        st.session_state["final_filename"] = _fname

    with st.expander("📖 Preview Report (Dify sections)", expanded=True):
        st.markdown(result_text)

    # ── Inputs used — a collapsed read-only summary of exactly what produced
    #    this report. The input form is hidden once a report exists, so this
    #    read-only snapshot is the only place the chosen values survive.
    with st.expander("📋 Inputs used for this report", expanded=False):
        _flag = lambda v: "✓" if v else "—"
        _g = lambda k, d="": ui.get(k, d)

        st.markdown("**Report parameters**")
        st.markdown(
            f"- Report type: `{_g('Report_type')}`\n"
            f"- Output language: `{_g('Output_language')}`\n"
            f"- Scope of report (SSO): `{_g('Scope_of_the_report')}`\n"
            f"- Period: `{_g('Period_start')}` → `{_g('Period_end')}`"
        )

        st.markdown("**Company / service**")
        st.markdown(
            f"- Company name: {_g('Company_name') or '—'}\n"
            f"- Short name: {_g('Co_short_name') or '—'}\n"
            f"- Industry: {_g('Industry') or '—'}\n"
            f"- System / service: {_g('System_or_service_name') or '—'}\n"
            f"- Service description: {_g('Service_description') or '—'}\n"
            f"- System (extra): {_g('System') or '—'}\n"
            f"- Systems function: {_g('Systems_function') or '—'}\n"
            f"- Domain: {_g('Domain') or '—'}\n"
            f"- Subservice org: {_g('Subservice_org') or '—'}\n"
            f"- Website: {_g('Co_website') or '—'}"
        )

        st.markdown("**Trust Service Criteria (SOC2)**")
        st.markdown(
            f"- Security {_flag(_g('is_Security'))} · "
            f"Availability {_flag(_g('is_Availability'))} · "
            f"Processing Integrity {_flag(_g('is_Processing_Integrity'))} · "
            f"Confidentiality {_flag(_g('is_Confidentiality'))} · "
            f"Privacy {_flag(_g('is_Privacy'))}"
        )

        st.markdown("**User Entity sections**")
        st.markdown(
            f"- CUEC (Complementary User Entity Controls): {_flag(_g('is_CUEC'))}\n"
            f"- UER (User Entity Responsibilities): {_flag(_g('is_UER'))}"
        )

        if tc.get("generate_complete"):
            st.markdown("**Complete-report (MA + AR) settings**")
            st.markdown(
                f"- Standard: `{tc.get('standard') or '—'}`\n"
                f"- Report signing date: {tc.get('report_date') or '—'}\n"
                f"- Signing city: {tc.get('signing_city') or '—'}\n"
                f"- AR addressee: {tc.get('addressee_choice') or '—'}\n"
                f"- CUEC identified: {_flag(tc.get('cuec_identified'))}\n"
                f"- SSO complementary controls identified: {_flag(tc.get('sso_cc_identified'))}\n"
                f"- Transaction processing wording: {_flag(tc.get('has_transaction_processing'))}\n"
                f"- Single user entity: {_flag(tc.get('single_user_entity'))}\n"
                f"- AI scope exclusion: {_flag(tc.get('has_ai_scope_exclusion'))}\n"
                f"- Other Information section: {_flag(tc.get('has_other_information'))}"
            )
            _mp = tc.get("ma_template_path")
            _ap = tc.get("ar_template_path")
            st.markdown(
                f"- MA template: `{os.path.basename(_mp) if _mp else '—'}`\n"
                f"- AR template: `{os.path.basename(_ap) if _ap else '—'}`"
            )

    final_bytes = st.session_state["final_bytes"]
    filename    = st.session_state["final_filename"]

    st.download_button(
        label="⬇ Download Report (.docx)",
        data=final_bytes,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary",
        use_container_width=True,
    )
